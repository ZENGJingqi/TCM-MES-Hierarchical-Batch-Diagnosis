from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


CN_TO_EN = {
    "批号": "batch_no",
    "生产日期": "production_date",
    "包衣片重(g)": "coated_tablet_weight_g",
    "崩解时限(min)": "disintegration_time_min",
    "成分含量(mg/片)": "active_content_mg_per_tablet",
    "需氧菌总数(cfu/g)": "total_aerobic_count_cfu_per_g",
    "霉菌及酵母菌(cfu/g)": "mold_yeast_count_cfu_per_g",
    "规格": "dosage_strength",
}

INDICATOR_LABELS = {
    "coated_tablet_weight_g": "Coated tablet weight (g)",
    "disintegration_time_min": "Disintegration time (min)",
    "active_content_mg_per_tablet": "Active content (mg/tablet)",
}


def prettify_strength(series: pd.Series) -> pd.Series:
    series = series.astype(str).str.strip().str.replace(r"\s+", "", regex=True)
    return series.str.replace(r"^([0-9]+(?:\.[0-9]+)?)g$", r"\1 g", regex=True)


def fmt_num(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}"


def build_data(input_path: Path) -> pd.DataFrame:
    df = pd.read_excel(input_path)
    df = df.rename(columns=CN_TO_EN)
    df["production_date"] = pd.to_datetime(df["production_date"])
    df["production_month"] = df["production_date"].dt.to_period("M").dt.to_timestamp()
    df["dosage_strength"] = prettify_strength(df["dosage_strength"])
    return df


def compute_spec_summary(df: pd.DataFrame) -> pd.DataFrame:
    grouped = df.groupby("dosage_strength", sort=False)
    rows = []
    for strength, tmp in grouped:
        rows.append(
            {
                "dosage_strength": strength,
                "records": int(len(tmp)),
                "batches": int(tmp["batch_no"].nunique()),
                "date_start": tmp["production_date"].min().date().isoformat(),
                "date_end": tmp["production_date"].max().date().isoformat(),
                "weight_mean": tmp["coated_tablet_weight_g"].mean(),
                "weight_sd": tmp["coated_tablet_weight_g"].std(),
                "weight_median": tmp["coated_tablet_weight_g"].median(),
                "weight_q1": tmp["coated_tablet_weight_g"].quantile(0.25),
                "weight_q3": tmp["coated_tablet_weight_g"].quantile(0.75),
                "weight_min": tmp["coated_tablet_weight_g"].min(),
                "weight_max": tmp["coated_tablet_weight_g"].max(),
                "dis_mean": tmp["disintegration_time_min"].mean(),
                "dis_sd": tmp["disintegration_time_min"].std(),
                "dis_median": tmp["disintegration_time_min"].median(),
                "dis_q1": tmp["disintegration_time_min"].quantile(0.25),
                "dis_q3": tmp["disintegration_time_min"].quantile(0.75),
                "dis_min": tmp["disintegration_time_min"].min(),
                "dis_max": tmp["disintegration_time_min"].max(),
                "act_mean": tmp["active_content_mg_per_tablet"].mean(),
                "act_sd": tmp["active_content_mg_per_tablet"].std(),
                "act_median": tmp["active_content_mg_per_tablet"].median(),
                "act_q1": tmp["active_content_mg_per_tablet"].quantile(0.25),
                "act_q3": tmp["active_content_mg_per_tablet"].quantile(0.75),
                "act_min": tmp["active_content_mg_per_tablet"].min(),
                "act_max": tmp["active_content_mg_per_tablet"].max(),
            }
        )
    return pd.DataFrame(rows)


def compute_monthly(df: pd.DataFrame) -> pd.DataFrame:
    monthly = (
        df.groupby(["dosage_strength", "production_month"], sort=True)
        .agg(
            n=("batch_no", "size"),
            weight_mean=("coated_tablet_weight_g", "mean"),
            dis_mean=("disintegration_time_min", "mean"),
            dis_sd=("disintegration_time_min", "std"),
            act_mean=("active_content_mg_per_tablet", "mean"),
            gt10_n=("disintegration_time_min", lambda x: int((x > 10).sum())),
            ge10_n=("disintegration_time_min", lambda x: int((x >= 10).sum())),
        )
        .reset_index()
    )
    monthly["gt10_pct"] = monthly["gt10_n"] / monthly["n"] * 100
    monthly["ge10_pct"] = monthly["ge10_n"] / monthly["n"] * 100
    return monthly


def build_core_metric_table(spec_summary: pd.DataFrame) -> pd.DataFrame:
    rows = []
    metric_map = {
        "coated_tablet_weight_g": ("weight_mean", "weight_sd", "weight_median", "weight_q1", "weight_q3", "weight_min", "weight_max"),
        "disintegration_time_min": ("dis_mean", "dis_sd", "dis_median", "dis_q1", "dis_q3", "dis_min", "dis_max"),
        "active_content_mg_per_tablet": ("act_mean", "act_sd", "act_median", "act_q1", "act_q3", "act_min", "act_max"),
    }
    for _, row in spec_summary.iterrows():
        for metric, cols in metric_map.items():
            rows.append(
                {
                    "Specification": row["dosage_strength"],
                    "Indicator": INDICATOR_LABELS[metric],
                    "n": int(row["records"]),
                    "Mean ± SD": f"{fmt_num(row[cols[0]])} ± {fmt_num(row[cols[1]])}",
                    "Median (Q1, Q3)": f"{fmt_num(row[cols[2]])} ({fmt_num(row[cols[3]])}, {fmt_num(row[cols[4]])})",
                    "Range": f"{fmt_num(row[cols[5]])} to {fmt_num(row[cols[6]])}",
                }
            )
    return pd.DataFrame(rows)


def build_threshold_table(monthly: pd.DataFrame) -> pd.DataFrame:
    monthly_08 = monthly[monthly["dosage_strength"] == "0.8 g"].copy()
    keep_months = ["2024-08", "2024-09", "2025-12", "2026-01", "2026-02"]
    monthly_08["month_str"] = monthly_08["production_month"].dt.strftime("%Y-%m")
    monthly_08 = monthly_08[monthly_08["month_str"].isin(keep_months)]
    monthly_08 = monthly_08[["month_str", "n", "dis_mean", "gt10_n", "gt10_pct", "ge10_n", "ge10_pct"]]
    monthly_08 = monthly_08.rename(
        columns={
            "month_str": "Period",
            "n": "n",
            "dis_mean": "Mean disintegration time (min)",
            "gt10_n": ">10 min (n)",
            "gt10_pct": ">10 min (%)",
            "ge10_n": ">=10 min (n)",
            "ge10_pct": ">=10 min (%)",
        }
    )
    monthly_08["Mean disintegration time (min)"] = monthly_08["Mean disintegration time (min)"].map(lambda x: fmt_num(x))
    monthly_08[">10 min (%)"] = monthly_08[">10 min (%)"].map(lambda x: f"{x:.2f}")
    monthly_08[">=10 min (%)"] = monthly_08[">=10 min (%)"].map(lambda x: f"{x:.2f}")
    return monthly_08


def add_table(document: Document, dataframe: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(dataframe.columns):
        hdr_cells[idx].text = str(col)
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def dataframe_to_markdown(dataframe: pd.DataFrame) -> str:
    headers = [str(col) for col in dataframe.columns]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for _, row in dataframe.iterrows():
        lines.append("| " + " | ".join(str(value) for value in row.tolist()) + " |")
    return "\n".join(lines)


def set_doc_style(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def build_text(
    df: pd.DataFrame,
    spec_summary: pd.DataFrame,
    monthly: pd.DataFrame,
) -> dict[str, str]:
    total_records = len(df)
    total_batches = df["batch_no"].nunique()
    date_start = df["production_date"].min().date().isoformat()
    date_end = df["production_date"].max().date().isoformat()

    s05 = spec_summary[spec_summary["dosage_strength"] == "0.5 g"].iloc[0]
    s08 = spec_summary[spec_summary["dosage_strength"] == "0.8 g"].iloc[0]

    m05 = monthly[monthly["dosage_strength"] == "0.5 g"].copy()
    m08 = monthly[monthly["dosage_strength"] == "0.8 g"].copy()

    d05_gt10 = int((df.loc[df["dosage_strength"] == "0.5 g", "disintegration_time_min"] > 10).sum())
    d05_ge10 = int((df.loc[df["dosage_strength"] == "0.5 g", "disintegration_time_min"] >= 10).sum())
    d08_gt10 = int((df.loc[df["dosage_strength"] == "0.8 g", "disintegration_time_min"] > 10).sum())
    d08_ge10 = int((df.loc[df["dosage_strength"] == "0.8 g", "disintegration_time_min"] >= 10).sum())

    p1 = (
        f"D2 成品理化数据共纳入 {total_records} 条记录，对应 {total_batches} 个唯一批次，覆盖 {date_start} 至 {date_end}。"
        f"其中，0.8 g 规格 {int(s08['records'])} 条，占全部记录的 {s08['records'] / total_records * 100:.2f}%；"
        f"0.5 g 规格 {int(s05['records'])} 条，占 {s05['records'] / total_records * 100:.2f}%。"
        "由于两种规格的时间覆盖范围和样本量明显不同，后续结果均按规格分别分析。"
    )

    p2 = (
        f"0.5 g 规格共 {int(s05['records'])} 条记录，覆盖 {s05['date_start']} 至 {s05['date_end']}。"
        f"片重整体较为稳定，均值为 {fmt_num(s05['weight_mean'])} g，标准差为 {fmt_num(s05['weight_sd'])} g，"
        f"月均值范围为 {fmt_num(m05['weight_mean'].min())} 至 {fmt_num(m05['weight_mean'].max())} g。"
        f"成分含量均值为 {fmt_num(s05['act_mean'])} mg/tablet，标准差为 {fmt_num(s05['act_sd'])} mg/tablet，"
        f"月均值范围为 {fmt_num(m05['act_mean'].min())} 至 {fmt_num(m05['act_mean'].max())} mg/tablet。"
        f"崩解时限均值为 {fmt_num(s05['dis_mean'])} min，标准差为 {fmt_num(s05['dis_sd'])} min，"
        f"月均值自 2025-01 的 {fmt_num(m05.iloc[0]['dis_mean'])} min 上升至 2026-02 的 {fmt_num(m05.iloc[-1]['dis_mean'])} min，"
        f"总体月均值范围为 {fmt_num(m05['dis_mean'].min())} 至 {fmt_num(m05['dis_mean'].max())} min。"
        f"在全部 0.5 g 记录中，仅 {d05_gt10} 条记录超过 10 min，{d05_ge10} 条记录达到或超过 10 min，"
        "说明该规格下崩解时限虽有一定抬升，但整体仍处于相对稳定状态。"
    )

    p3 = (
        f"0.8 g 规格共 {int(s08['records'])} 条记录，覆盖 {s08['date_start']} 至 {s08['date_end']}。"
        f"片重均值为 {fmt_num(s08['weight_mean'])} g，标准差为 {fmt_num(s08['weight_sd'])} g，"
        f"月均值仅在 {fmt_num(m08['weight_mean'].min())} 至 {fmt_num(m08['weight_mean'].max())} g 之间波动，整体稳定性较好。"
        f"成分含量均值为 {fmt_num(s08['act_mean'])} mg/tablet，标准差为 {fmt_num(s08['act_sd'])} mg/tablet，"
        f"月均值范围为 {fmt_num(m08['act_mean'].min())} 至 {fmt_num(m08['act_mean'].max())} mg/tablet，"
        "存在一定阶段性波动，但未表现出与片重同方向的大幅失稳。"
    )

    p4 = (
        f"与片重和成分含量相比，0.8 g 规格的崩解时限表现出更突出的阶段性变化。"
        f"该指标整体均值为 {fmt_num(s08['dis_mean'])} min，标准差为 {fmt_num(s08['dis_sd'])} min，"
        f"总体范围为 {fmt_num(s08['dis_min'])} 至 {fmt_num(s08['dis_max'])} min。"
        f"从时间序列看，2024-08 和 2024-09 的月均值分别为 {fmt_num(m08.iloc[0]['dis_mean'])} 和 {fmt_num(m08.iloc[1]['dis_mean'])} min，"
        "处于较高水平；随后自 2024-10 至 2025-11 基本维持在 7 至 9 min 区间。"
        f"自 2025-12 起，崩解时限再次抬升，月均值为 {fmt_num(m08.loc[m08['production_month'] == pd.Timestamp('2025-12-01'), 'dis_mean'].iloc[0])} min，"
        f"到 2026-01 和 2026-02 分别升至 {fmt_num(m08.loc[m08['production_month'] == pd.Timestamp('2026-01-01'), 'dis_mean'].iloc[0])} 和 {fmt_num(m08.loc[m08['production_month'] == pd.Timestamp('2026-02-01'), 'dis_mean'].iloc[0])} min。"
    )

    p5 = (
        f"按阈值观察，0.8 g 规格中共有 {d08_gt10} 条记录超过 10 min，占该规格全部记录的 {d08_gt10 / s08['records'] * 100:.2f}%；"
        f"若按大于等于 10 min 统计，则共有 {d08_ge10} 条记录，占 {d08_ge10 / s08['records'] * 100:.2f}%。"
        f"其中，2025-12、2026-01 和 2026-02 超过 10 min 的比例分别为 "
        f"{m08.loc[m08['production_month'] == pd.Timestamp('2025-12-01'), 'gt10_pct'].iloc[0]:.2f}%、"
        f"{m08.loc[m08['production_month'] == pd.Timestamp('2026-01-01'), 'gt10_pct'].iloc[0]:.2f}% 和 "
        f"{m08.loc[m08['production_month'] == pd.Timestamp('2026-02-01'), 'gt10_pct'].iloc[0]:.2f}%。"
        "这一结果表明，D2 中最值得进一步追踪的问题并不是片重或成分含量的整体漂移，而是 0.8 g 规格崩解时限在特定阶段出现的集中性升高。"
    )

    brief = (
        "D2 的核心结果可以概括为：两种规格的片重总体稳定，成分含量存在一定波动但未表现出持续性失控；"
        "真正需要优先追踪的异常信号集中在 0.8 g 规格的崩解时限。"
        "该指标在 2024-08 至 2024-09 曾处于高位，随后长期维持在相对较低水平，但在 2025-12 至 2026-02 再次明显升高，"
        "其中 2026-01 和 2026-02 超过 10 min 的记录占比分别达到 72.84% 和 84.72%。"
    )

    return {
        "overview": p1,
        "spec_05": p2,
        "spec_08": p3,
        "issue_1": p4,
        "issue_2": p5,
        "brief": brief,
    }


def write_markdown(
    output_path: Path,
    core_metric_table: pd.DataFrame,
    threshold_table: pd.DataFrame,
    paragraphs: dict[str, str],
) -> None:
    lines = [
        "# D2 正式结果分析与文字整理",
        "",
        "## 一、结果概述",
        "",
        paragraphs["brief"],
        "",
        "## 二、三项核心理化指标描述统计",
        "",
        dataframe_to_markdown(core_metric_table),
        "",
        "## 三、0.8 g 崩解时限阈值汇总",
        "",
        dataframe_to_markdown(threshold_table),
        "",
        "## 四、可直接用于论文中文稿的结果段落",
        "",
        paragraphs["overview"],
        "",
        paragraphs["spec_05"],
        "",
        paragraphs["spec_08"],
        "",
        paragraphs["issue_1"],
        "",
        paragraphs["issue_2"],
        "",
        "## 五、对应图件",
        "",
        "- 0.5 g 片重：01_0p5g_overview_coated_tablet_weight",
        "- 0.5 g 崩解时限：02_0p5g_overview_disintegration_time",
        "- 0.5 g 成分含量：03_0p5g_overview_active_content",
        "- 0.8 g 片重：04_0p8g_overview_coated_tablet_weight",
        "- 0.8 g 崩解时限：05_0p8g_overview_disintegration_time",
        "- 0.8 g 成分含量：06_0p8g_overview_active_content",
        "",
    ]
    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_docx(
    output_path: Path,
    core_metric_table: pd.DataFrame,
    threshold_table: pd.DataFrame,
    paragraphs: dict[str, str],
) -> None:
    document = Document()
    set_doc_style(document)

    title = document.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("D2 正式结果分析与文字整理")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    document.add_heading("一、结果概述", level=1)
    document.add_paragraph(paragraphs["brief"])

    document.add_heading("二、三项核心理化指标描述统计", level=1)
    add_table(document, core_metric_table)

    document.add_heading("三、0.8 g 崩解时限阈值汇总", level=1)
    add_table(document, threshold_table)

    document.add_heading("四、可直接用于论文中文稿的结果段落", level=1)
    for key in ["overview", "spec_05", "spec_08", "issue_1", "issue_2"]:
        document.add_paragraph(paragraphs[key])

    document.add_heading("五、对应图件", level=1)
    for item in [
        "0.5 g 片重：01_0p5g_overview_coated_tablet_weight",
        "0.5 g 崩解时限：02_0p5g_overview_disintegration_time",
        "0.5 g 成分含量：03_0p5g_overview_active_content",
        "0.8 g 片重：04_0p8g_overview_coated_tablet_weight",
        "0.8 g 崩解时限：05_0p8g_overview_disintegration_time",
        "0.8 g 成分含量：06_0p8g_overview_active_content",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.save(output_path)


def main() -> None:
    script_path = Path(__file__).resolve()
    project_dir = script_path.parents[1]
    root_dir = project_dir.parent
    docs_dir = project_dir / "docs"
    docs_dir.mkdir(parents=True, exist_ok=True)

    input_path = root_dir / "定稿数据_中文" / "D2_成品理化数据_定稿.xlsx"
    df = build_data(input_path)
    spec_summary = compute_spec_summary(df)
    monthly = compute_monthly(df)
    core_metric_table = build_core_metric_table(spec_summary)
    threshold_table = build_threshold_table(monthly)
    paragraphs = build_text(df, spec_summary, monthly)

    md_path = docs_dir / "D2_正式结果分析与文字整理.md"
    docx_path = docs_dir / "D2_正式结果分析与文字整理.docx"

    write_markdown(md_path, core_metric_table, threshold_table, paragraphs)
    write_docx(docx_path, core_metric_table, threshold_table, paragraphs)

    print(f"Saved markdown: {md_path}")
    print(f"Saved docx: {docx_path}")


if __name__ == "__main__":
    main()
