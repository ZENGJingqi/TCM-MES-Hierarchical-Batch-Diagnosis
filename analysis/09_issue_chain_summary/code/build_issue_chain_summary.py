from __future__ import annotations

from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(r"D:\博士后文件\论文撰写\华润江中大数据分析")
OUT_DIR = BASE_DIR / "成品崩解问题驱动链整合分析"
DOCS_DIR = OUT_DIR / "docs"
TABLES_DIR = OUT_DIR / "tables"


def find_one(pattern: str) -> Path:
    matches = list(BASE_DIR.rglob(pattern))
    if not matches:
        raise FileNotFoundError(pattern)
    matches.sort(key=lambda p: len(str(p)))
    return matches[0]


def load_excel(path: Path) -> dict[str, pd.DataFrame]:
    xl = pd.ExcelFile(path)
    return {sheet: xl.parse(sheet) for sheet in xl.sheet_names}


def format_num(value: float | int, digits: int = 2) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, int):
        return str(value)
    return f"{value:.{digits}f}"


def get_row_value(df: pd.DataFrame, key_col: str, key: str, value_col: str = "value"):
    row = df.loc[df[key_col] == key]
    if row.empty:
        raise KeyError(f"{key} not found")
    return row.iloc[0][value_col]


def bold_run(paragraph, text: str, bold: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    return run


def set_document_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in doc.styles:
            doc.styles[name].font.name = "Times New Roman"
            doc.styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_table(doc: Document, dataframe: pd.DataFrame, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bold_run(p, title, bold=True)

    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, col in enumerate(dataframe.columns):
        header_cells[idx].text = str(col)

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row.tolist()):
            cells[idx].text = "" if pd.isna(value) else str(value)
    doc.add_paragraph("")


def write_markdown(path: Path, sections: Iterable[str]) -> None:
    path.write_text("\n\n".join(sections).strip() + "\n", encoding="utf-8")


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    TABLES_DIR.mkdir(parents=True, exist_ok=True)

    d2_final_path = find_one("D2_成品理化数据_定稿.xlsx")
    d2_summary_path = find_one("D2_descriptive_summary.xlsx")
    d4_path = find_one("D4_finished_disintegration_association.xlsx")
    d6_path = find_one("D6_finished_disintegration_association.xlsx")
    d7_path = find_one("D7_finished_disintegration_association.xlsx")

    d2_final = pd.read_excel(d2_final_path)
    d2_summary = load_excel(d2_summary_path)
    d4 = load_excel(d4_path)
    d6 = load_excel(d6_path)
    d7 = load_excel(d7_path)

    d2_spec_col = d2_final.columns[7]
    d2_date_col = d2_final.columns[1]
    d2_dis_col = d2_final.columns[3]
    d2_batch_col = d2_final.columns[0]

    d2_08 = d2_final.loc[
        d2_final[d2_spec_col].astype(str).str.replace(" ", "", regex=False).eq("0.8g")
    ].copy()
    d2_08[d2_date_col] = pd.to_datetime(d2_08[d2_date_col])
    d2_08["production_month"] = d2_08[d2_date_col].dt.to_period("M").astype(str)
    d2_08["issue_flag"] = d2_08[d2_dis_col] > 10
    d2_month = (
        d2_08.groupby("production_month")
        .agg(
            records=(d2_dis_col, "size"),
            issue_n=("issue_flag", "sum"),
        )
        .reset_index()
    )
    d2_month["issue_rate_pct"] = (d2_month["issue_n"] / d2_month["records"] * 100).round(2)
    peak_months = d2_month.sort_values("issue_rate_pct", ascending=False).head(4)

    overall_d2 = d2_summary["overall_overview"]
    spec_d2 = d2_summary["spec_overview"]
    core_d2 = d2_summary["core_metric_summary"]
    d2_08_spec = spec_d2.loc[spec_d2["dosage_strength"].astype(str).str.contains("0.8")]

    d4_pair = d4["pairing_overview"]
    d4_group = d4["group_comparison"]
    d4_uni = d4["univariate_ratio_assoc"]
    d4_glm = d4["multivariable_glm"]

    d6_pair = d6["pairing_overview"]
    d6_group = d6["group_comparison"]
    d6_uni = d6["issue_ratio_assoc"]
    d6_glm = d6["logistic_glm"]
    d6_origin = d6["origin_pattern_summary"]
    d6_origin_single = d6["origin_single_result"]

    d7_pair = d7["pairing_overview"]
    d7_group = d7["group_comparison"]
    d7_uni = d7["issue_ratio_assoc"]
    d7_glm = d7["logistic_glm"]
    d7_supplier = d7["supplier_summary"]
    d7_supplier_test = d7["supplier_fisher_result"]

    endpoint_summary = pd.DataFrame(
        [
            {
                "层级": "成品终点",
                "分析对象": "0.8 g 成品批次",
                "记录数": len(d2_08),
                "唯一批次": d2_08[d2_batch_col].nunique(),
                "问题定义": "崩解时限 > 10 min",
                "问题记录数": int(d2_08["issue_flag"].sum()),
                "问题比例(%)": round(d2_08["issue_flag"].mean() * 100, 2),
                "时间范围": f"{d2_08[d2_date_col].min().date()} 至 {d2_08[d2_date_col].max().date()}",
            }
        ]
    )

    linkage_summary = pd.DataFrame(
        [
            {
                "上游主体": "浸膏粉",
                "上游可配对实体数": int(get_row_value(d4_pair, "item", "D2-D3-D4 matched extract batches")),
                "可配对成品批次数": int(get_row_value(d4_pair, "item", "D2-D3-D4 matched finished batches")),
                "问题上游实体数": int(get_row_value(d4_pair, "item", "Matched extract batches with at least one linked batch >10 min")),
                "问题上游实体比例(%)": round(
                    int(get_row_value(d4_pair, "item", "Matched extract batches with at least one linked batch >10 min"))
                    / int(get_row_value(d4_pair, "item", "D2-D3-D4 matched extract batches"))
                    * 100,
                    2,
                ),
                "配对说明": "D2 -> D3 -> D4，主分析单位为浸膏粉批次",
            },
            {
                "上游主体": "陈皮",
                "上游可配对实体数": int(get_row_value(d6_pair, "item", "Linked chenpi batches")),
                "可配对成品批次数": int(get_row_value(d6_pair, "item", "Linked 0.8g finished-product batches")),
                "问题上游实体数": int(str(get_row_value(d6_pair, "item", "Chenpi batches with >=1 linked batch >10 min")).split(" ")[0]),
                "问题上游实体比例(%)": 15.62,
                "配对说明": "D6 -> D5 -> D3 -> D2，主分析单位为陈皮/浸膏粉链路批次",
            },
            {
                "上游主体": "山药粉",
                "上游可配对实体数": int(get_row_value(d7_pair, "item", "Linked yam-powder batches")),
                "可配对成品批次数": int(get_row_value(d7_pair, "item", "Linked 0.8g finished-product batches")),
                "问题上游实体数": int(str(get_row_value(d7_pair, "item", "Yam-powder batches with >=1 linked batch >10 min")).split(" ")[0]),
                "问题上游实体比例(%)": 25.19,
                "配对说明": "D7 -> D3 -> D2，主分析单位为山药粉批次",
            },
        ]
    )

    evidence_summary = pd.DataFrame(
        [
            {
                "层级": "浸膏粉",
                "关键变量": "总灰分(%)",
                "方向": "升高 -> 崩解问题风险升高",
                "单变量相关": "rho = 0.3995, P < 0.001",
                "组间差异": "5.20 vs 4.70, P < 0.001",
                "多变量结果": "OR = 9.4848, 95%CI 5.9966 to 15.0022, P < 0.001",
                "判定": "主正向信号",
            },
            {
                "层级": "浸膏粉",
                "关键变量": "浸出物(%)",
                "方向": "降低 -> 崩解问题风险升高",
                "单变量相关": "rho = -0.3135, P < 0.001",
                "组间差异": "42.00 vs 43.70, P < 0.001",
                "多变量结果": "OR = 0.1364, 95%CI 0.0755 to 0.2466, P < 0.001",
                "判定": "主负向信号",
            },
            {
                "层级": "陈皮",
                "关键变量": "橙皮苷(%)",
                "方向": "降低 -> 崩解问题风险升高",
                "单变量相关": "rho = -0.4995, P < 0.001",
                "组间差异": "5.6 vs 6.6, P < 0.001",
                "多变量结果": "OR = 0.0800, 95%CI 0.0281 to 0.2282, P < 0.001",
                "判定": "最强陈皮信号",
            },
            {
                "层级": "陈皮",
                "关键变量": "水分(%)",
                "方向": "降低 -> 崩解问题风险升高",
                "单变量相关": "rho = -0.4338, P < 0.001",
                "组间差异": "2.1 vs 8.0, P < 0.001",
                "多变量结果": "OR = 0.3794, 95%CI 0.1498 to 0.9605, P = 0.041",
                "判定": "稳定负向信号",
            },
            {
                "层级": "山药粉",
                "关键变量": "挑选出不合格品占比(%)",
                "方向": "升高 -> 崩解问题风险升高",
                "单变量相关": "rho = 0.3751, P < 0.001",
                "组间差异": "0.0284 vs 0.0207, P < 0.001",
                "多变量结果": "OR = 1.8212, 95%CI 1.1229 to 2.9536, P = 0.015",
                "判定": "主正向信号",
            },
            {
                "层级": "山药粉",
                "关键变量": "120目细度均值(%)",
                "方向": "降低 -> 崩解问题风险升高",
                "单变量相关": "rho = -0.3278, P < 0.001",
                "组间差异": "99.6750 vs 99.7333, P < 0.001",
                "多变量结果": "OR = 0.3366, 95%CI 0.1967 to 0.5763, P < 0.001",
                "判定": "主负向信号",
            },
            {
                "层级": "山药粉",
                "关键变量": "120目细度SD",
                "方向": "增大 -> 崩解问题风险升高",
                "单变量相关": "rho = 0.3711, P < 0.001",
                "组间差异": "0.1054 vs 0.0577, P < 0.001",
                "多变量结果": "未进入最终GLM，但方向稳定",
                "判定": "辅助波动信号",
            },
        ]
    )

    stratified_summary = pd.DataFrame(
        [
            {
                "层级": "陈皮来源",
                "分层结果": "浙江 vs 江西宜春",
                "主要结果": "27.38% vs 2.76%",
                "统计学": "OR = 13.1333, 95%CI 4.2402 to 54.4919, P < 0.001",
                "说明": "当前 D6 只有产地字段，没有显式厂家字段",
            },
            {
                "层级": "山药粉供应商",
                "分层结果": "江西樟树天齐堂 / 洛阳康鑫 / 亳州中信 / 河南尚华堂 / Missing",
                "主要结果": "100% / 100% / 55% / 0% / 0%",
                "统计学": "Fisher-Freeman-Halton P < 0.001",
                "说明": "供应商分布极不均衡，小样本组需谨慎解释",
            },
        ]
    )

    peak_months_display = peak_months.assign(
        描述=lambda x: x["production_month"]
        + ": "
        + x["issue_n"].astype(str)
        + "/"
        + x["records"].astype(str)
        + " ("
        + x["issue_rate_pct"].map(lambda v: f"{v:.2f}%")
        + ")"
    )["描述"].tolist()

    sections = [
        (
            "一、终点定义与总体概况",
            f"本轮问题驱动分析以 0.8 g 成品批次的崩解时限 > 10 min 作为统一终点。当前定稿数据中，"
            f"0.8 g 成品共 {len(d2_08)} 条记录、{d2_08[d2_batch_col].nunique()} 个唯一批次，"
            f"其中崩解异常记录 {int(d2_08['issue_flag'].sum())} 条，占 {d2_08['issue_flag'].mean()*100:.2f}%。"
            f"异常并非均匀分布，而是呈阶段性集中，问题率最高的月份包括 "
            + "；".join(peak_months_display)
            + "。整体上看，成品崩解异常在早期上市阶段和 2026 年 1 至 2 月两个时间窗口最为突出，因此适合作为后续上游追溯分析的明确终点。"
        ),
        (
            "二、浸膏粉层证据",
            f"浸膏粉层通过 D2 -> D3 -> D4 共匹配到 {int(get_row_value(d4_pair, 'item', 'D2-D3-D4 matched finished batches'))} 个成品批次和 "
            f"{int(get_row_value(d4_pair, 'item', 'D2-D3-D4 matched extract batches'))} 个浸膏粉批次，"
            f"其中 {int(get_row_value(d4_pair, 'item', 'Matched extract batches with at least one linked batch >10 min'))} 个浸膏粉批次至少关联 1 个异常成品批次。"
            "浸膏粉层最稳定的异常模式为“总灰分升高、浸出物降低”。总灰分与问题批比例呈正相关，"
            f"Spearman rho = {d4_uni.loc[d4_uni['Indicator']=='Total ash (%)', 'Spearman rho with linked >10 min ratio'].iloc[0]:.4f}, "
            f"P {d4_uni.loc[d4_uni['Indicator']=='Total ash (%)', 'P value'].iloc[0]}；"
            f"多变量模型中 OR = {d4_glm.loc[d4_glm['Indicator']=='Total ash (%)', 'OR per 1 SD increase'].iloc[0]:.4f}。"
            "浸出物与问题批比例呈负相关，且在多变量模型中仍保持显著保护方向。"
            "因此，浸膏粉层可以把成品崩解问题进一步收敛到“灰分/浸出物”这一组上游质量信号。"
        ),
        (
            "三、陈皮层证据",
            f"陈皮层通过 D6 -> D5 -> D3 -> D2 共匹配到 {int(get_row_value(d6_pair, 'item', 'Linked chenpi batches'))} 个陈皮批次、"
            f"{int(get_row_value(d6_pair, 'item', 'Linked extract-powder batches'))} 个浸膏粉批次和 "
            f"{int(get_row_value(d6_pair, 'item', 'Linked 0.8g finished-product batches'))} 个成品批次。"
            "陈皮层最强信号来自橙皮苷和水分，两者均表现为数值降低时，后续成品崩解异常风险上升。"
            f"其中橙皮苷与问题批比例的相关性为 Spearman rho = {d6_uni.loc[d6_uni['Indicator']=='Chenpi hesperidin (%)', 'Spearman rho with linked >10 min ratio'].iloc[0]:.4f}，"
            f"多变量 OR = {d6_glm.loc[d6_glm['Indicator']=='Chenpi hesperidin (%)', 'Odds ratio'].iloc[0]:.4f}；"
            f"水分与问题批比例的相关性为 Spearman rho = {d6_uni.loc[d6_uni['Indicator']=='Chenpi moisture (%)', 'Spearman rho with linked >10 min ratio'].iloc[0]:.4f}，"
            f"多变量 OR = {d6_glm.loc[d6_glm['Indicator']=='Chenpi moisture (%)', 'Odds ratio'].iloc[0]:.4f}。"
            "此外，来源分层结果显示浙江来源对应的后续问题率显著高于江西宜春。"
            "需要强调的是，D6 当前只有产地信息，没有显式厂家字段，因此这一层应解释为“来源差异”，而不是“厂家差异”。"
        ),
        (
            "四、山药粉层证据",
            f"山药粉层通过 D7 -> D3 -> D2 共匹配到 {int(get_row_value(d7_pair, 'item', 'Linked yam-powder batches'))} 个山药粉批次和 "
            f"{int(get_row_value(d7_pair, 'item', 'Linked 0.8g finished-product batches'))} 个成品批次，"
            f"其中 {str(get_row_value(d7_pair, 'item', 'Yam-powder batches with >=1 linked batch >10 min'))} 的山药粉批次关联到了异常成品批次。"
            "山药粉层的主信号主要集中在分选损耗和细度，而不是得率或物料平衡。"
            f"挑选出不合格品占比与问题批比例正相关，Spearman rho = {d7_uni.loc[d7_uni['Indicator']=='Rejected material rate (%)', 'Spearman rho with linked >10 min ratio'].iloc[0]:.4f}；"
            f"120 目细度均值与问题批比例负相关，Spearman rho = {d7_uni.loc[d7_uni['Indicator']=='Through 120-mesh mean (%)', 'Spearman rho with linked >10 min ratio'].iloc[0]:.4f}。"
            "多变量模型中，这两类信号仍然保留。供应商分层也显示出明显差异，但由于部分类别样本量极小，当前更适合作为风险分层线索，而不宜直接写成稳固结论。"
        ),
        (
            "五、综合判断",
            "把终点与三层上游证据合在一起，可以得到一条比较清楚的问题驱动链："
            "成品端的崩解异常并不是孤立现象，而是能够被上游多层级质量属性持续追溯。"
            "其中，浸膏粉层负责把问题收敛到“总灰分升高/浸出物降低”；"
            "陈皮层进一步提示“低橙皮苷、低水分、浙江来源”与该模式相关；"
            "山药粉层则提示“分选损耗升高、120 目细度下降且波动增大”同样与终点异常相关。"
            "三层证据并不是相互替代，而是从不同主体共同指向同一个成品终点，符合当前论文想强调的“问题驱动、分层传递、批次级追溯”框架。"
        ),
        (
            "六、下一步建议",
            "下一步不建议继续做更松散的单变量扩展，而应进入联合建模。"
            "最有价值的工作是：1）基于已筛出的关键变量建立一张跨层级批次证据矩阵；"
            "2）用成品崩解异常作为终点，做跨层级联合分类模型；"
            "3）将结果整理为论文主文中的“问题链图 + 主结果表”，把其余结果降为补充材料。"
        ),
    ]

    narrative_sections = ["# 成品崩解问题驱动链整合分析（定稿）"] + [
        f"## {title}\n{body}" for title, body in sections
    ]

    markdown_path = DOCS_DIR / "成品崩解问题驱动链整合分析_定稿.md"
    write_markdown(markdown_path, narrative_sections)

    doc = Document()
    set_document_style(doc)
    doc.add_heading("成品崩解问题驱动链整合分析（定稿）", 0)

    intro = doc.add_paragraph()
    intro.add_run(
        "本稿以 0.8 g 成品批次崩解时限 > 10 min 作为统一终点，"
        "将成品、浸膏粉、陈皮和山药粉四层结果收敛为一条可用于论文写作的问题驱动证据链。"
    )

    doc.add_heading("一、终点与配对规模", level=1)
    add_table(doc, endpoint_summary, "表1 终点定义")
    add_table(doc, linkage_summary, "表2 各层级与成品终点的配对规模")

    p = doc.add_paragraph()
    p.add_run("0.8 g 成品的崩解异常记录为 ").bold = False
    p.add_run(f"{int(d2_08['issue_flag'].sum())}/{len(d2_08)} ({d2_08['issue_flag'].mean()*100:.2f}%)").bold = True
    p.add_run("。问题率最高的月份包括：")
    p.add_run("；".join(peak_months_display)).bold = True
    p.add_run("。")

    doc.add_heading("二、关键证据摘要", level=1)
    add_table(doc, evidence_summary, "表3 各层级关键显著因素摘要")
    add_table(doc, stratified_summary, "表4 来源/供应商分层结果")

    for title, body in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)

    docx_path = DOCS_DIR / "成品崩解问题驱动链整合分析_定稿.docx"
    doc.save(docx_path)

    xlsx_path = TABLES_DIR / "成品崩解问题驱动链整合证据表.xlsx"
    with pd.ExcelWriter(xlsx_path) as writer:
        endpoint_summary.to_excel(writer, sheet_name="endpoint_summary", index=False)
        linkage_summary.to_excel(writer, sheet_name="linkage_summary", index=False)
        evidence_summary.to_excel(writer, sheet_name="evidence_summary", index=False)
        stratified_summary.to_excel(writer, sheet_name="stratified_summary", index=False)
        peak_months.to_excel(writer, sheet_name="endpoint_peak_months", index=False)

    print(f"Wrote: {docx_path}")
    print(f"Wrote: {markdown_path}")
    print(f"Wrote: {xlsx_path}")


if __name__ == "__main__":
    main()
