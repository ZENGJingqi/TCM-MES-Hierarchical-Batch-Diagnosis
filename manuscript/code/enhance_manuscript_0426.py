from __future__ import annotations

import os
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path.cwd()


def find_file(filename: str) -> Path:
    for dirpath, _, filenames in os.walk(ROOT):
        if filename in filenames:
            return Path(dirpath) / filename
    raise FileNotFoundError(filename)


def find_manuscript_dir() -> Path:
    for dirpath, _, filenames in os.walk(ROOT):
        if "论文初稿-0426.docx" in filenames:
            return Path(dirpath)
    raise FileNotFoundError("论文初稿-0426.docx")


MANUSCRIPT_DIR = find_manuscript_dir()
SOURCE_DOC = MANUSCRIPT_DIR / "论文初稿-0426.docx"
ASCII_SOURCE = ROOT / "manuscript_working_0426.docx"
OUTPUT_DOC = MANUSCRIPT_DIR / "论文初稿-0426_完善版.docx"


def copy_source_to_ascii() -> None:
    # python-docx can be unreliable with Chinese paths in this Windows shell, so work through an ASCII copy.
    shutil.copyfile(SOURCE_DOC, ASCII_SOURCE)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(9)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)

    for style_name in ["Normal", "Body Text"]:
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            style.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.bold = True
        if style_name == "Heading 1":
            style.font.size = Pt(14)
        elif style_name == "Heading 2":
            style.font.size = Pt(12)
        else:
            style.font.size = Pt(11)


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(9)
    r.font.italic = True


def add_picture(doc: Document, filename: str, caption: str, width_cm: float = 15.5) -> None:
    path = find_file(filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_dataset_table(doc: Document) -> None:
    add_heading(doc, "Table 1. Overview of real-world batch-level manufacturing data", 2)
    headers = [
        "Data entity",
        "Source",
        "Records",
        "Batch-level linkage",
        "Role in this study",
    ]
    rows = [
        [
            "Jianwei Xiaoshi tablet quality testing data",
            "Quality-control testing",
            "3,728",
            "Linked to tablet MES records by finished-product batch identifier; 908 records were linked.",
            "Defines the finished-product quality outcome and disintegration issue.",
        ],
        [
            "Jianwei Xiaoshi tablet MES production records",
            "MES production records",
            "1,243",
            "Links tablet batches to Jianwei Xiaoshi extract-powder and Chinese yam powder input batches.",
            "Provides the batch-linkage backbone for upstream diagnosis.",
        ],
        [
            "Jianwei Xiaoshi extract-powder quality testing data",
            "Quality-control testing",
            "618",
            "Linked to tablet MES records by extract-powder batch identifier; 933 MES records were linked.",
            "Represents upstream extract-powder quality information.",
        ],
        [
            "Chenpi quality testing data",
            "Quality-control testing",
            "44",
            "Connected to Jianwei Xiaoshi extract-powder batches through Chenpi batch traceability; 41 Chenpi testing batches were linked.",
            "Represents a key herbal-material quality layer.",
        ],
        [
            "Chinese yam powder MES production records",
            "MES production records",
            "342",
            "Linked to tablet MES records by Chinese yam powder batch identifier; 1,159 MES records were linked.",
            "Represents the process-material production layer.",
        ],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9E7EA")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)


def add_evidence_table(doc: Document) -> None:
    add_heading(doc, "Table 2. Layer-wise evidence for finished-product disintegration diagnosis", 2)
    headers = ["Layer", "Linked samples", "Main signal", "Statistical evidence", "Interpretation"]
    rows = [
        [
            "Finished-product quality",
            "3,375 records for the 0.8 g specification",
            "Disintegration time >10 min was concentrated in specific periods.",
            "412/3,375 records exceeded 10 min; 2026-01 and 2026-02 reached 72.84% and 84.72%, respectively.",
            "The 0.8 g disintegration time was selected as the issue-driven endpoint.",
        ],
        [
            "Jianwei Xiaoshi extract powder",
            "804 linked finished-product records and 241 extract-powder batches",
            "Higher total ash and lower extractives were associated with longer disintegration.",
            "Total ash: Spearman rho = 0.311, P <0.001; OR per SD = 9.485. Extractives: rho = -0.210, P = 0.001; OR per SD = 0.136.",
            "Extract-powder quality narrowed the issue to a total-ash/extractives pattern.",
        ],
        [
            "Chenpi",
            "873 linked finished-product records, 242 extract-powder batches, and 32 Chenpi batches",
            "Lower Chenpi hesperidin and moisture were associated with a higher subsequent issue rate.",
            "Hesperidin: rho = -0.500, P <0.001. Moisture: rho = -0.434, P <0.001.",
            "Chenpi attributes suggested upstream raw-material contribution to the extract-powder pattern.",
        ],
        [
            "Chinese yam powder",
            "850 linked finished-product records and 131 Chinese yam powder batches",
            "Higher rejected-material rate and lower 120-mesh fineness were associated with the issue endpoint.",
            "Rejected-material rate: rho = 0.375. 120-mesh fineness mean: rho = -0.328.",
            "Chinese yam powder provided a parallel process-material risk signal.",
        ],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "E9E1D7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)


def clear_doc(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def build_doc() -> None:
    copy_source_to_ascii()
    doc = Document(str(ASCII_SOURCE))
    clear_doc(doc)
    style_doc(doc)

    title = (
        "Tracing finished-product quality issues to upstream material and process risks in traditional "
        "Chinese medicine manufacturing using MES-enabled hierarchical batch diagnosis"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(15)
    r.font.bold = True

    add_heading(doc, "Abstract", 1)
    add_paragraph(
        doc,
        "Traditional Chinese medicine manufacturing involves multiple herbal materials, intermediate extract preparations, "
        "process-material production, and finished-product manufacturing steps. These features create complex batch-to-batch "
        "relationships that are difficult to interpret using isolated quality-control records. This study developed an "
        "MES-enabled hierarchical batch diagnosis strategy to trace finished-product quality issues to upstream material "
        "and process risks using real-world Jianwei Xiaoshi tablet manufacturing data. A total of 5,975 valid batch-level "
        "records from August 2024 to March 2026 were curated, including Jianwei Xiaoshi tablet quality testing data, "
        "tablet MES production records, Jianwei Xiaoshi extract-powder testing data, Chenpi testing data, and Chinese yam "
        "powder MES records. Finished-product disintegration time in the 0.8 g specification was selected as the issue-driven "
        "endpoint. The analysis identified temporally concentrated disintegration increases and traced the endpoint to "
        "upstream signals involving extract-powder total ash and extractives, Chenpi hesperidin and moisture, and Chinese yam "
        "powder rejected-material rate and 120-mesh fineness. These findings demonstrate that real-world MES data can be "
        "used as a batch-linkage backbone for hierarchical diagnosis of finished-product quality issues in traditional "
        "Chinese medicine manufacturing."
    )
    add_paragraph(doc, "Keywords: traditional Chinese medicine manufacturing; MES; batch traceability; disintegration time; quality diagnosis; real-world manufacturing data.")

    add_heading(doc, "Introduction", 1)
    add_paragraph(
        doc,
        "Quality variation in traditional Chinese medicine products is shaped by the combined influence of herbal raw materials, "
        "intermediate extract preparation, process-material manufacturing, and finished-product production. Compared with many "
        "single-component pharmaceutical products, traditional Chinese medicine manufacturing often involves multiple herbal "
        "materials, multiple origins, batch mixing, and intermediate materials that are carried forward into many finished-product "
        "batches. Therefore, a finished-product quality issue is rarely interpretable from a single endpoint measurement alone."
    )
    add_paragraph(
        doc,
        "Manufacturing execution systems (MES) provide routine batch-level production records and material-input relationships. "
        "However, MES data are often used for retrospective documentation rather than as an analytical backbone for quality diagnosis. "
        "For real-world manufacturing data, a key challenge is to transform fragmented quality-control and production records into "
        "a coherent batch-level evidence chain that can support issue-driven investigation."
    )
    add_paragraph(
        doc,
        "Here, we propose an issue-driven hierarchical batch diagnosis framework and apply it to real-world Jianwei Xiaoshi tablet "
        "manufacturing data. The framework starts from a finished-product quality issue, reconstructs batch linkages through MES and "
        "traceability information, evaluates upstream material and process-material signals, and integrates evidence across layers. "
        "The goal is not only to describe quality variation, but to demonstrate a practical framework for tracing finished-product "
        "quality issues to upstream material and process risks."
    )

    add_heading(doc, "Materials and methods", 1)
    add_heading(doc, "Real-world manufacturing dataset", 2)
    add_paragraph(
        doc,
        "This study used real-world manufacturing data from Jianwei Xiaoshi tablets and integrated 5,975 valid batch-level records "
        "from quality-control testing and MES production records between August 2024 and March 2026. The dataset covered five data "
        "entities: 3,728 Jianwei Xiaoshi tablet quality-testing records, 1,243 Jianwei Xiaoshi tablet MES production records, "
        "618 Jianwei Xiaoshi extract-powder quality-testing records, 44 Chenpi quality-testing records, and 342 Chinese yam powder "
        "MES production records. These records were connected through batch identifiers, including finished-product batch identifiers, "
        "Jianwei Xiaoshi extract-powder batch identifiers, Chenpi batch identifiers, and Chinese yam powder batch identifiers."
    )
    add_dataset_table(doc)
    add_picture(
        doc,
        "Figure_1_dataset_overview_and_batch_linkage.png",
        "Figure 1. Overview of the real-world batch-level manufacturing dataset and batch-linkage structure. The figure shows the five data entities, record counts, and observed linkage coverage across Jianwei Xiaoshi tablet quality testing, tablet MES records, Jianwei Xiaoshi extract powder, Chenpi, and Chinese yam powder.",
        width_cm=16.5,
    )

    add_heading(doc, "Definition of the finished-product issue endpoint", 2)
    add_paragraph(
        doc,
        "Finished-product disintegration time was used as the issue-driven endpoint because preliminary profiling showed that tablet "
        "weight and active content were comparatively stable, whereas disintegration time displayed clear temporal concentration in "
        "the 0.8 g specification. A threshold of >10 min was used to define the disintegration issue for downstream hierarchical "
        "diagnosis."
    )
    add_heading(doc, "Batch linkage and hierarchical analysis", 2)
    add_paragraph(
        doc,
        "Finished-product quality records were first linked to tablet MES records using finished-product batch identifiers. MES "
        "records were then used to link finished-product batches to Jianwei Xiaoshi extract-powder batches and Chinese yam powder "
        "batches. Chenpi quality records were mapped to the extract-powder layer through Chenpi batch traceability information. "
        "Associations between upstream variables and finished-product disintegration were evaluated using rank correlation, group "
        "comparisons based on the >10 min endpoint, and multivariable models where appropriate."
    )

    add_heading(doc, "Results", 1)
    add_heading(doc, "The real-world dataset supported hierarchical batch linkage", 2)
    add_paragraph(
        doc,
        "Among the 3,728 Jianwei Xiaoshi tablet quality-testing records, 908 records were linked to tablet MES production records. "
        "Within the MES records, 933 records were linked to Jianwei Xiaoshi extract-powder testing data and 1,159 records were linked "
        "to Chinese yam powder MES records. End-to-end linkage enabled 804 tablet quality-testing records to be traced through MES to "
        "Jianwei Xiaoshi extract-powder quality data and 839 records to Chinese yam powder MES records. In addition, 41 Chenpi testing "
        "batches were connected to the extract-powder layer through traceability information."
    )

    add_heading(doc, "Finished-product disintegration time defined the primary issue", 2)
    add_paragraph(
        doc,
        "The finished-product quality dataset included 3,728 records, of which 3,375 belonged to the 0.8 g specification and 353 to "
        "the 0.5 g specification. In the 0.8 g specification, coated tablet weight was stable, with a mean of 0.7977 g and a standard "
        "deviation of 0.0031 g. Active content also showed only moderate stage-specific variation. In contrast, disintegration time "
        "showed a clear temporal pattern, with 412 of 3,375 records (12.21%) exceeding 10 min."
    )
    add_paragraph(
        doc,
        "The >10 min disintegration issue was concentrated in several production periods. The issue rate reached 94.44% in 2024-08 "
        "and 82.95% in 2024-09, decreased during the intervening period, and increased again in 2026-01 and 2026-02, reaching 72.84% "
        "and 84.72%, respectively. This temporal concentration supported the use of 0.8 g disintegration time >10 min as the primary "
        "endpoint for upstream diagnosis."
    )
    add_picture(
        doc,
        "05_0p8g_overview_disintegration_time.png",
        "Figure 2. Temporal and distributional overview of disintegration time for the 0.8 g Jianwei Xiaoshi tablet specification. The 10 min reference line marks the issue threshold used for hierarchical diagnosis.",
        width_cm=15.5,
    )

    add_heading(doc, "Jianwei Xiaoshi extract-powder quality was associated with finished-product disintegration", 2)
    add_paragraph(
        doc,
        "Using extract-powder batch identifiers in tablet MES records, 804 finished-product records were linked to 241 Jianwei Xiaoshi "
        "extract-powder batches. All linked finished-product records belonged to the 0.8 g specification. Among these linked records, "
        "79 records (9.83%) had disintegration time >10 min, and 26 extract-powder batches were linked to at least one such finished-product "
        "record."
    )
    add_paragraph(
        doc,
        "At the extract-powder batch level, total ash was positively associated with mean finished-product disintegration time "
        "(Spearman rho = 0.311, P <0.001), whereas extractives were negatively associated with disintegration time (rho = -0.210, "
        "P = 0.001). In multivariable models, total ash retained the strongest positive signal, with an odds ratio of 9.485 per "
        "standard-deviation increase for the >10 min endpoint, while extractives showed a protective direction with an odds ratio "
        "of 0.136 per standard-deviation increase."
    )
    add_picture(
        doc,
        "02_d4_finished_disintegration_scatter_total_ash.png",
        "Figure 3. Association between Jianwei Xiaoshi extract-powder total ash and linked finished-product disintegration time.",
        width_cm=13.5,
    )
    add_picture(
        doc,
        "03_d4_finished_disintegration_scatter_extract.png",
        "Figure 4. Association between Jianwei Xiaoshi extract-powder extractives and linked finished-product disintegration time.",
        width_cm=13.5,
    )

    add_heading(doc, "Chenpi quality contributed an upstream raw-material signal", 2)
    add_paragraph(
        doc,
        "Chenpi testing records were connected to the extract-powder layer through batch traceability information. After batch cleaning, "
        "41 Chenpi testing batches could be linked to 585 extract-powder batches, and 32 Chenpi batches were further linked to 811 "
        "finished-product batches through the Chenpi-extract powder-tablet chain. Chenpi attributes were associated with both extract-powder "
        "quality and the downstream disintegration endpoint."
    )
    add_paragraph(
        doc,
        "Chenpi hesperidin showed the strongest positive correlation with extract-powder extractives (Spearman rho = 0.500, P <0.001), "
        "whereas Chenpi moisture was negatively correlated with extract-powder total ash (rho = -0.400, P <0.001). When mapped to "
        "finished-product disintegration, lower Chenpi hesperidin and moisture were associated with a higher linked issue rate "
        "(rho = -0.500 and -0.434, respectively; both P <0.001). Source stratification further suggested a difference between Chenpi "
        "origins, although this result should be interpreted as a source-level pattern rather than a manufacturer-level conclusion."
    )
    add_picture(
        doc,
        "01_chenpi_extract_correlation_heatmap.png",
        "Figure 5. Correlation heatmap linking Chenpi quality attributes with Jianwei Xiaoshi extract-powder quality attributes.",
        width_cm=13.5,
    )

    add_heading(doc, "Chinese yam powder MES records provided a parallel process-material signal", 2)
    add_paragraph(
        doc,
        "Chinese yam powder MES records were linked to finished-product quality records through tablet MES material-input information. "
        "This linkage yielded 850 finished-product records and 131 Chinese yam powder batches, of which 33 batches (25.19%) were linked "
        "to at least one finished-product record with disintegration time >10 min."
    )
    add_paragraph(
        doc,
        "The main Chinese yam powder signals were concentrated in rejected-material rate and fineness. Rejected-material rate was "
        "positively correlated with the issue rate of linked finished-product batches (Spearman rho = 0.375), while the mean proportion "
        "passing through 120 mesh was negatively correlated with the issue rate (rho = -0.328). These findings suggest that Chinese yam "
        "powder production records contributed a process-material evidence layer that was complementary to the extract-powder and Chenpi layers."
    )
    add_picture(
        doc,
        "01_d7_finished_group_rejected_material_rate.png",
        "Figure 6. Chinese yam powder rejected-material rate grouped by whether the linked finished-product batches included a disintegration issue.",
        width_cm=13.5,
    )
    add_picture(
        doc,
        "02_d7_finished_group_through_120_mesh_mean.png",
        "Figure 7. Chinese yam powder 120-mesh fineness grouped by whether the linked finished-product batches included a disintegration issue.",
        width_cm=13.5,
    )

    add_heading(doc, "Integrated hierarchical diagnosis", 2)
    add_paragraph(
        doc,
        "Integrating the finished-product endpoint with upstream evidence produced a coherent issue-driven chain. The finished-product "
        "signal was a temporally concentrated increase in 0.8 g disintegration time. The extract-powder layer narrowed the issue to "
        "a total-ash and extractives pattern. The Chenpi layer provided upstream raw-material evidence, especially hesperidin and moisture. "
        "The Chinese yam powder layer provided an additional process-material signal involving rejected-material rate and 120-mesh fineness. "
        "These layers did not replace one another; instead, they provided complementary batch-level evidence for tracing the same finished-product issue."
    )
    add_evidence_table(doc)

    add_heading(doc, "Discussion", 1)
    add_paragraph(
        doc,
        "This study demonstrates that routine MES and quality-control records can be reorganized into a hierarchical evidence chain for "
        "finished-product quality diagnosis. The key methodological contribution is the issue-driven linkage strategy: the analysis begins "
        "with a finished-product quality issue, reconstructs batch relationships through MES, and then evaluates upstream material and "
        "process-material signals. This structure is particularly relevant for traditional Chinese medicine manufacturing, where multiple "
        "herbal materials, intermediate extracts, and process-material records jointly influence finished-product quality."
    )
    add_paragraph(
        doc,
        "The current findings should be interpreted as real-world batch-level diagnostic evidence rather than controlled experimental causality. "
        "Batch linkage was incomplete because the available records covered different time ranges and data systems. Some upstream signals, "
        "particularly source-related Chenpi differences and supplier-related Chinese yam powder differences, require cautious interpretation "
        "because subgroup sizes may be uneven. Nevertheless, the concordance of evidence across finished-product, extract-powder, Chenpi, and "
        "Chinese yam powder layers supports the value of MES-enabled hierarchical diagnosis."
    )
    add_paragraph(
        doc,
        "Future work should extend this framework into a formal cross-layer evidence matrix and interpretable graph-based model. Such a model "
        "could represent finished-product batches, material batches, and process records as heterogeneous nodes connected by batch-use edges, "
        "with attention-based weighting used to prioritize upstream risk evidence. This would move the framework from retrospective diagnosis "
        "toward decision-support for real-world pharmaceutical manufacturing."
    )

    add_heading(doc, "Data availability", 1)
    add_paragraph(
        doc,
        "A curated English data overview and variable dictionary was prepared as a supplementary workbook: "
        "real_world_manufacturing_dataset_overview_and_dictionary.xlsx."
    )
    add_heading(doc, "Supplementary materials", 1)
    add_paragraph(doc, "Supplementary Table S1. Real-world manufacturing dataset overview and variable dictionary.")
    add_paragraph(doc, "Supplementary Table S2. Layer-wise batch-linkage and evidence summary for finished-product disintegration diagnosis.")
    add_paragraph(doc, "Supplementary Figures. Additional descriptive and association plots for finished-product, extract-powder, Chenpi, and Chinese yam powder datasets.")

    doc.save(str(ROOT / "manuscript_enhanced_ascii.docx"))
    shutil.copyfile(ROOT / "manuscript_enhanced_ascii.docx", OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    build_doc()
