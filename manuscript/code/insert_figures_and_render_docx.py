from pathlib import Path
import subprocess

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_DIR = Path.cwd()
WRITE_DIR = PROJECT_DIR / "\u8bba\u6587\u5199\u4f5c"
FIG_DIR = PROJECT_DIR / "\u8bba\u6587\u4e3b\u56fe_\u6846\u67b6\u4e0e\u8bc1\u636e\u94fe" / "figures"
SOURCE_DOCX = WRITE_DIR / "\u8bba\u6587\u6846\u67b6_\u6b63\u6587\u4e0e\u652f\u6491\u6750\u6599_\u6295\u7a3f\u7ea7\u8bbe\u8ba1\u7a3f.docx"
OUT_DOCX = WRITE_DIR / "\u8bba\u6587\u6846\u67b6_\u6b63\u6587\u4e0e\u652f\u6491\u6750\u6599_\u6295\u7a3f\u7ea7\u8bbe\u8ba1\u7a3f_\u542b\u4e3b\u56fe.docx"
RENDER_DIR = WRITE_DIR / "rendered_preview"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")

FIGURES = [
    (
        "Figure 1. MES-enabled issue-driven hierarchical batch diagnosis framework.",
        FIG_DIR / "Figure_framework_issue_driven_hierarchical_batch_diagnosis.png",
        "\u8be5\u56fe\u7528\u4e8e\u5c55\u793a\u5168\u6587\u6846\u67b6\uff1a\u4ece\u6210\u54c1\u5d29\u89e3\u5f02\u5e38\u8bc6\u522b\uff0c\u5230\u6279\u6b21\u8ffd\u6eaf\u3001\u9010\u5c42\u7b5b\u9009\u3001\u8054\u5408\u5efa\u6a21\u3001\u65f6\u95f4\u7a97\u53e3\u6821\u6b63\u3001\u56e0\u679c\u542f\u53d1\u5f0f\u8def\u5f84\u5206\u89e3\u548c\u56fe\u8bc1\u636e\u8bc4\u5206\u3002",
    ),
    (
        "Figure 6. Evidence chain from finished-product disintegration issue to upstream risk-priority output.",
        FIG_DIR / "Figure_evidence_chain_finished_issue_to_upstream_risk_priority.png",
        "\u8be5\u56fe\u7528\u4e8e\u5c55\u793a\u6700\u7ec8\u8bc1\u636e\u94fe\uff1a\u6210\u54c1\u5d29\u89e3\u95ee\u9898\u901a\u8fc7 MES\u3001\u6d78\u818f\u7c89\u3001\u9648\u76ae\u548c\u5c71\u836f\u7c89\u6279\u6b21\u5173\u8054\uff0c\u8fdb\u5165 graph-based batch evidence scoring\uff0c\u8f93\u51fa\u4e0a\u6e38\u6279\u6b21\u98ce\u9669\u4f18\u5148\u7ea7\u3002",
    ),
]


def set_run_font(run, size=None, bold=None, italic=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_figure_block(doc, caption, image_path, note, page_break_before=False):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.page_break_before = page_break_before
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Cm(25.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_run_font(r, size=10.5, bold=True)

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_note.paragraph_format.space_after = Pt(4)
    nr = p_note.add_run(note)
    set_run_font(nr, size=9.5)


def render_pdf_preview(pdf_path):
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    for old in RENDER_DIR.glob(f"{pdf_path.stem}_page_*.png"):
        old.unlink()
    doc = fitz.open(str(pdf_path))
    pages = []
    for idx in range(min(len(doc), 12)):
        page = doc.load_page(idx)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
        out = RENDER_DIR / f"{pdf_path.stem}_page_{idx + 1:02d}.png"
        pix.save(str(out))
        pages.append(out)
    doc.close()
    return pages


def main():
    doc = Document(str(SOURCE_DOCX))

    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    intro = doc.add_paragraph("\u6b63\u6587\u4e3b\u56fe\u9884\u89c8\uff08\u6a2a\u7248\uff0c\u5df2\u63d2\u5165\uff09", style="Heading 1")
    set_run_font(intro.runs[0], size=16, bold=True)
    p = doc.add_paragraph(
        "\u4ee5\u4e0b\u4e24\u5f20\u56fe\u5df2\u63d2\u5165 Word \u6b63\u6587\u6587\u6863\u4e2d\uff0c\u4f5c\u4e3a\u6295\u7a3f\u7ea7\u4e3b\u56fe\u8bbe\u8ba1\u7684\u5f53\u524d\u7248\u672c\u9884\u89c8\u3002"
    )
    set_run_font(p.runs[0], size=11)

    for idx, (caption, image_path, note) in enumerate(FIGURES):
        add_figure_block(doc, caption, image_path, note, page_break_before=idx > 0)

    doc.save(str(OUT_DOCX))

    pdf_path = None
    if SOFFICE.exists():
        pdf_path = OUT_DOCX.with_suffix(".pdf")
        if pdf_path.exists():
            pdf_path.unlink()
        subprocess.run(
            [
                str(SOFFICE),
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(WRITE_DIR),
                str(OUT_DOCX),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        pages = render_pdf_preview(pdf_path)
    else:
        pages = []

    print("DOCX:", OUT_DOCX)
    if pdf_path:
        print("PDF:", pdf_path)
    for p in pages:
        print("PREVIEW:", p)


if __name__ == "__main__":
    main()
