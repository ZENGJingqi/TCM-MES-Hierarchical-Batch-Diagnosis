from pathlib import Path
import subprocess

import fitz
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_DIR = Path.cwd()
WRITE_DIR = PROJECT_DIR / "论文写作"
OUT_DOCX = WRITE_DIR / "投稿级论文正文与支撑材料设计稿_含图表.docx"
OUT_PDF = OUT_DOCX.with_suffix(".pdf")
RENDER_DIR = WRITE_DIR / "rendered_preview"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")


FIGS = {
    "framework": PROJECT_DIR / "论文主图_框架与证据链" / "figures" / "Figure_framework_issue_driven_hierarchical_batch_diagnosis.png",
    "data": PROJECT_DIR / "数据概况图_英文" / "figures" / "Figure_1_dataset_overview_and_batch_linkage.png",
    "d2_disintegration": PROJECT_DIR / "成品问题驱动_D2描述" / "figures" / "05_0p8g_overview_disintegration_time.png",
    "d2_weight": PROJECT_DIR / "成品问题驱动_D2描述" / "figures" / "08_0p8g_disintegration_group_coated_tablet_weight.png",
    "d2_content": PROJECT_DIR / "成品问题驱动_D2描述" / "figures" / "09_0p8g_disintegration_group_active_content.png",
    "d3_rank": PROJECT_DIR / "D3_成品MES与成品崩解关联分析" / "figures" / "01_d3_mes_feature_screening_rank.png",
    "d4_rank": PROJECT_DIR / "D4_浸膏粉与成品崩解关联分析" / "figures" / "00_d4_extract_feature_screening_rank.png",
    "d6_rank": PROJECT_DIR / "D6_陈皮与成品崩解关联分析" / "figures" / "00_d6_chenpi_feature_screening_rank.png",
    "d7_rank": PROJECT_DIR / "D7_山药粉与成品崩解关联分析" / "figures" / "00_d7_yam_powder_feature_screening_rank.png",
    "d3_key": PROJECT_DIR / "D3_成品MES与成品崩解关联分析" / "figures" / "05_d3_key_mes_features_vs_disintegration_time.png",
    "d4_key": PROJECT_DIR / "D4_浸膏粉与成品崩解关联分析" / "figures" / "08_d4_key_extract_features_vs_disintegration_time.png",
    "d6_key": PROJECT_DIR / "D6_陈皮与成品崩解关联分析" / "figures" / "08_d6_key_chenpi_features_vs_disintegration_issue_ratio.png",
    "d7_key": PROJECT_DIR / "D7_山药粉与成品崩解关联分析" / "figures" / "08_d7_key_yam_powder_features_vs_disintegration_issue_ratio.png",
    "model_auc": PROJECT_DIR / "联合建模_成品崩解问题驱动" / "figures" / "05_publication_model_auc_prauc.png",
    "model_brier": PROJECT_DIR / "联合建模_成品崩解问题驱动" / "figures" / "06_publication_model_brier_score.png",
    "model_coef": PROJECT_DIR / "联合建模_成品崩解问题驱动" / "figures" / "07_publication_clean_core_coefficients.png",
    "time_series": PROJECT_DIR / "时序分析_成品崩解问题" / "figures" / "01_finished_disintegration_time_series.png",
    "monthly_issue": PROJECT_DIR / "时序分析_成品崩解问题" / "figures" / "02_monthly_disintegration_issue_rate.png",
    "time_model": PROJECT_DIR / "时序分析_成品崩解问题" / "figures" / "04_time_adjusted_model_performance.png",
    "causal_exposure": PROJECT_DIR / "因果建模_成品崩解问题" / "figures" / "02_upstream_issue_rate_vs_abnormal_window_exposure.png",
    "causal_bootstrap": PROJECT_DIR / "因果建模_成品崩解问题" / "figures" / "04_bootstrap_batch_level_effects.png",
    "causal_chenpi": PROJECT_DIR / "因果建模_成品崩解问题" / "figures" / "05_chenpi_source_path_attenuation.png",
    "evidence_chain": PROJECT_DIR / "论文主图_框架与证据链" / "figures" / "Figure_evidence_chain_finished_issue_to_upstream_risk_priority.png",
    "graph_top": PROJECT_DIR / "批次证据图模型_成品崩解问题" / "figures" / "01_batch_evidence_graph_top_risk_nodes.png",
    "graph_landscape": PROJECT_DIR / "批次证据图模型_成品崩解问题" / "figures" / "02_time_adjusted_batch_evidence_landscape.png",
    "graph_distribution": PROJECT_DIR / "批次证据图模型_成品崩解问题" / "figures" / "03_layerwise_graph_evidence_score_distribution.png",
}


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def style_doc(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    styles = doc.styles
    for name in ["Normal", "Body Text"]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(31, 31, 31)


def add_p(doc, text="", size=10.5, bold=False, italic=False, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=10.2)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = hdr[i].paragraphs[0].add_run(h)
        set_run_font(run, size=9.2, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cells[i].paragraphs[0].add_run(str(val))
            set_run_font(run, size=8.7)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return table


def set_landscape(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.0)


def set_portrait(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)


def add_figure(doc, key, caption, note=None, width_cm=25.6, break_before=True):
    if break_before:
        doc.add_page_break()
    path = FIGS[key]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        with Image.open(path) as img:
            img_w, img_h = img.size
        max_h_cm = 14.3
        if (img_w / img_h) >= (width_cm / max_h_cm):
            p.add_run().add_picture(str(path), width=Cm(width_cm))
        else:
            p.add_run().add_picture(str(path), height=Cm(max_h_cm))
    else:
        run = p.add_run(f"[Missing figure: {path}]")
        set_run_font(run, size=10, bold=True, color=(180, 0, 0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=9.8, bold=True)
    if note:
        add_p(doc, note, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)


def add_two_figures(doc, left_key, left_caption, right_key, right_caption, break_before=True):
    if break_before:
        doc.add_page_break()
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (key, caption) in enumerate([(left_key, left_caption), (right_key, right_caption)]):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        path = FIGS[key]
        if path.exists():
            with Image.open(path) as img:
                img_w, img_h = img.size
            max_w_cm = 12.4
            max_h_cm = 11.5
            if (img_w / img_h) >= (max_w_cm / max_h_cm):
                p.add_run().add_picture(str(path), width=Cm(max_w_cm))
            else:
                p.add_run().add_picture(str(path), height=Cm(max_h_cm))
        c = table.cell(1, i).paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.add_run(caption)
        set_run_font(run, size=8.6, bold=True)


def render_pdf_preview(pdf_path):
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    for old in RENDER_DIR.glob(f"{pdf_path.stem}_page_*.png"):
        old.unlink()
    doc = fitz.open(str(pdf_path))
    pages = []
    for idx in range(min(len(doc), 30)):
        page = doc.load_page(idx)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.45, 1.45), alpha=False)
        out = RENDER_DIR / f"{pdf_path.stem}_page_{idx + 1:02d}.png"
        pix.save(str(out))
        pages.append(out)
    doc.close()
    return pages


def build_doc():
    doc = Document()
    style_doc(doc)

    title = "投稿级论文正文与支撑材料设计稿"
    add_p(doc, title, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_p(
        doc,
        "Tracing finished-product quality issues to upstream material and process risks in traditional Chinese medicine manufacturing using MES-enabled hierarchical batch diagnosis",
        size=11.5,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    add_p(
        doc,
        "用途：本文件用于固定论文正文大纲、Results/Methods/Discussion 写作骨架、正文主图主表和支撑材料体系。正文聚焦框架与关键证据，补充材料承接完整筛选、完整模型和完整变量说明。",
        size=10.5,
    )

    doc.add_heading("一、论文定位与核心问题", level=1)
    add_p(
        doc,
        "本文不是单纯报告若干显著相关指标，而是提出并验证一个面向真实世界中药制造的 MES-enabled hierarchical batch diagnosis 框架。核心问题是：当成品质量异常被发现后，如何把分散的 QC、MES 和批次追溯数据组织成可解释、可复核、可用于现场排查的上游物料与过程风险证据链。",
    )
    add_bullets(
        doc,
        [
            "问题入口：以 0.8 g 健胃消食片成品崩解时限 >10 min 作为 sentinel finished-product issue。",
            "数据基础：整合 2024 年 8 月至 2026 年 3 月真实世界质量控制检测和 MES 生产记录，共 5,975 条有效批次级记录。",
            "方法主线：issue identification -> batch linkage -> layer-wise screening -> joint modeling -> temporal adjustment -> causal-informed decomposition -> graph evidence scoring。",
            "输出形式：不是直接给出单一原因，而是输出上游批次 risk-priority ranking，用于现场聚焦排查和 CAPA 线索生成。",
        ],
    )

    doc.add_heading("二、正文大纲设计", level=1)
    outline_rows = [
        ["Introduction", "真实世界中药制造存在多饮片、多中间体、多工序和多对多批次流转，常规 QC 难以解释成品异常的上游来源。", "提出 MES-enabled hierarchical batch diagnosis 的必要性。"],
        ["Results 1", "真实世界多层级制造数据集与批次关联结构。", "Figure 1; Table 1"],
        ["Results 2", "成品崩解异常作为问题驱动入口。", "Figure 2; Supplementary Table S3"],
        ["Results 3", "逐层筛选候选过程与物料质量信号。", "Figure 3; Supplementary Table S4"],
        ["Results 4", "多层级联合建模证明 MES 和层级批次信息的诊断增益。", "Figure 4; Table 2"],
        ["Results 5", "时序窗口校正与因果启发式路径分解。", "Figure 5; Supplementary Table S6"],
        ["Results 6", "图证据评分将多层级证据转化为上游风险优先级。", "Figure 6; Table 3"],
        ["Discussion", "强调框架价值、MES 转化价值、中药制造批次结构特点、边界和未来验证。", "不作强因果、不作前瞻预测过度声明。"],
        ["Methods", "数据标准化、批次关联、异常定义、逐层筛选、联合建模、时序/因果分析、图证据评分。", "按可复现分析流程写。"],
    ]
    add_table(doc, ["章节", "正文内容", "图表/边界"], outline_rows, widths=[3.0, 9.0, 5.0])

    doc.add_heading("三、正文 Results 设计", level=1)
    result_sections = [
        (
            "Results 1. Real-world hierarchical manufacturing dataset and batch linkage",
            "本研究整合 5,975 条有效批次级记录，覆盖成品理化检测、成品 MES 生产记录、健胃消食片浸膏粉检测、陈皮检测和山药粉 MES 生产记录。该部分需要明确数据不是单表分析，而是多层级批次网络：成品质量记录可向 MES、浸膏粉、陈皮和山药粉生产过程逐层追溯。",
        ),
        (
            "Results 2. Finished-product disintegration as the sentinel quality issue",
            "成品端应突出 0.8 g 崩解时限 >10 min 的阶段性异常。片重和含量作为辅助质量属性，用于说明异常并非简单由片重偏移解释；全文主问题聚焦崩解时限。",
        ),
        (
            "Results 3. Layer-wise candidate evidence screening",
            "在成品 MES、浸膏粉、陈皮和山药粉层面分别进行候选指标筛选。正文只展示筛选秩图和每层最关键证据；完整箱线图、散点图、P 值、FDR 和样本量进入补充材料。",
        ),
        (
            "Results 4. Hierarchical joint modeling",
            "模型部分用于证明多层级数据的诊断增益。关键写法是：Baseline AUC 约 0.668，加入 MES 后 AUC 约 0.978，完整核心层级模型 AUC 约 0.981。时间外推表现较差，应解释为 retrospective diagnosis，而不是前瞻预测。",
        ),
        (
            "Results 5. Temporal adjustment and causal-informed decomposition",
            "时序分析说明异常窗口是强系统状态信号。因果启发式分析用于避免把陈皮来源或山药粉供应商直接写成原因；重点是路径衰减、时间窗口控制和批次层级稳健性。",
        ),
        (
            "Results 6. Graph-based batch evidence scoring",
            "图证据评分是论文落地模块。它将 time-adjusted residual、graph propagation、trace support 和 quality/process deviation 合成为 risk-priority score，输出可供现场排查的上游批次优先级。",
        ),
    ]
    for heading, text in result_sections:
        doc.add_heading(heading, level=2)
        add_p(doc, text)

    doc.add_heading("四、正文主表设计", level=1)
    table1 = [
        ["Finished-product quality testing", "3,728", "Finished-product batch identifier", "Defines disintegration issue and finished-product quality endpoints"],
        ["Tablet MES production records", "1,243", "Finished-product batch identifier; extract-powder and yam-powder batch identifiers", "Connects finished-product issue to production-process records"],
        ["Jianwei Xiaoshi extract-powder testing", "618", "Extract-powder batch identifier", "Intermediate quality bridge between raw material and tablet manufacturing"],
        ["Chenpi quality testing", "44", "Chenpi batch identifier", "Raw-material quality evidence linked through extract-powder traceability"],
        ["Chinese yam powder MES production records", "342", "Chinese yam powder batch identifier", "Process-material production evidence linked through tablet MES"],
    ]
    add_p(doc, "Table 1. Real-world manufacturing dataset and batch-linkage summary.", bold=True)
    add_table(doc, ["Data entity", "Records", "Primary linkage key", "Role in diagnosis"], table1, widths=[4.8, 2.1, 5.3, 5.5])

    table2 = [
        ["Baseline finished-product quality model", "Finished-product attributes only", "AUC ~0.668", "Limited diagnostic ability"],
        ["Baseline + tablet MES", "Finished quality + MES process parameters", "AUC ~0.978", "Largest gain; MES is central diagnostic layer"],
        ["Full core hierarchy", "Finished quality + MES + upstream quality/process evidence", "AUC ~0.981", "Marginal gain over MES but improves traceability and interpretation"],
        ["Time-split validation", "Training/testing separated by production time", "Performance unstable", "Supports retrospective diagnosis rather than direct prospective prediction"],
    ]
    add_p(doc, "Table 2. Hierarchical diagnostic model performance and interpretation.", bold=True)
    add_table(doc, ["Model", "Data layers", "Main result", "Interpretation"], table2, widths=[4.7, 5.4, 3.5, 4.7])

    table3 = [
        ["Extract powder", "25098043", "2/4", "83.6", "High-priority intermediate batch for focused investigation"],
        ["Chinese yam powder", "25106008", "1/2", "67.2", "High-priority process-material batch"],
        ["Chenpi", "2509801", "3/45", "31.7", "Raw-material batch with trace-supported downstream issue signal"],
    ]
    add_p(doc, "Table 3. Top-ranked upstream batches from graph-based evidence scoring.", bold=True)
    add_table(doc, ["Layer", "Batch", "Issue/linked records", "Risk-priority score", "Interpretation"], table3, widths=[3.5, 3.2, 3.5, 3.2, 5.5])

    doc.add_heading("五、Materials and Methods 写作骨架", level=1)
    methods = [
        ("Data curation and standardization", "说明中文原始数据如何标准化为 analysis-ready English datasets；批号转字符串；重复批号处理；多值字段保留原始序列并以均值列作为主分析变量；D7 缺失过程字段按均值插补。"),
        ("Batch linkage construction", "说明成品质量、成品 MES、浸膏粉、陈皮和山药粉之间如何通过批号建立层级关联，并区分 linked 和 traced 的定义。"),
        ("Definition of sentinel issue", "说明 0.8 g 成品崩解时限 >10 min 的定义依据，强调该阈值用于问题驱动诊断，不等同于法规失败或因果终点。"),
        ("Layer-wise statistical screening", "连续变量用 Spearman 相关；组间比较用 Wilcoxon 秩和检验；多重检验用 BH-FDR；上游批次按 linked issue ratio 或是否 linked >10 min batch 分析。"),
        ("Hierarchical joint modeling", "比较 Baseline、Baseline + MES、Full hierarchy 等模型；报告 AUC、PR-AUC、Brier score、重复交叉验证和时间切分验证。"),
        ("Temporal-window analysis", "按生产月份识别异常窗口，评估月度 issue rate、季节敏感性和异常窗口暴露。"),
        ("Causal-informed path decomposition", "采用路径衰减、时间窗口调整和 bootstrap batch-level effects，强调 causal-informed 而非 randomized causal inference。"),
        ("Graph-based batch evidence scoring", "构建 batch-evidence graph，并计算 time-adjusted residual、graph propagation、trace support、quality/process deviation 和 risk-priority score。"),
    ]
    add_table(doc, ["Method module", "Required description"], methods, widths=[5.0, 12.5])

    doc.add_heading("六、Discussion 写作骨架", level=1)
    add_bullets(
        doc,
        [
            "第一层：框架价值。本文解决的是成品终检异常与上游物料/过程风险之间证据链断裂的问题。",
            "第二层：中药制造特点。多饮片、多中间体、多批次流转和多对多追溯关系，使层级批次诊断比单表建模更有价值。",
            "第三层：MES 价值。MES 从记录系统转化为质量诊断和现场排查的证据系统。",
            "第四层：边界。当前为单产品、单企业、回顾性数据；不能声称特定产地或供应商直接导致异常；不能声称模型可直接前瞻预测。",
            "第五层：未来验证。需要在更多产品、更多时间窗口和实际 CAPA 场景中验证 risk-priority output 的实用性。",
        ],
    )

    doc.add_heading("七、支撑材料体系", level=1)
    supp_figs = [
        ["Supplementary Figure S1", "Finished-product quality attributes beyond disintegration", "片重、含量、崩解时限完整描述图。"],
        ["Supplementary Figure S2", "Finished-product MES feature screening", "成品 MES 所有有效过程参数筛选和关键参数关系图。"],
        ["Supplementary Figure S3", "Extract-powder quality feature screening", "浸膏粉指标与成品崩解关系。"],
        ["Supplementary Figure S4", "Chenpi quality and source-related analyses", "陈皮质量属性、来源模式和下游问题关联。"],
        ["Supplementary Figure S5", "Chinese yam powder MES feature screening", "山药粉 MES 指标与成品崩解关系。"],
        ["Supplementary Figure S6", "Elastic-net coefficients and SHAP importance", "联合模型变量贡献。"],
        ["Supplementary Figure S7", "Temporal and causal sensitivity analyses", "月度、季节、异常窗口、路径衰减和 bootstrap。"],
        ["Supplementary Figure S8", "Graph evidence score distribution and landscape", "图证据评分分布和 time-adjusted evidence landscape。"],
    ]
    add_p(doc, "Supplementary figures.", bold=True)
    add_table(doc, ["No.", "Title", "Content"], supp_figs, widths=[4.3, 6.2, 7.1])

    supp_tables = [
        ["Supplementary Table S1", "Data dictionary and variable-level completeness", "所有英文定稿数据表的字段说明、类型、缺失、范围。"],
        ["Supplementary Table S2", "Batch linkage and traceability summary", "各层级 linked/traced 批次数、记录数和关联比例。"],
        ["Supplementary Table S3", "Finished-product quality descriptive statistics", "成品理化描述和 >10 min 分组比较。"],
        ["Supplementary Table S4", "Layer-wise feature screening results", "所有有效指标的 P、FDR、效应方向和样本量。"],
        ["Supplementary Table S5", "Full hierarchical model results", "全部模型性能、交叉验证、时间拆分验证、核心变量。"],
        ["Supplementary Table S6", "Temporal and causal-informed analyses", "异常窗口、季节性、路径衰减、bootstrap batch-level effects。"],
        ["Supplementary Table S7", "Graph-based batch evidence scoring results", "所有上游批次 graph evidence score、risk-priority score 和解释标签。"],
        ["Supplementary Data 1", "Analysis-ready English dataset", "定稿英文数据文件。"],
        ["Supplementary Data 2", "Analysis code", "R 分析代码和图表生成代码。"],
    ]
    add_p(doc, "Supplementary tables and data.", bold=True)
    add_table(doc, ["No.", "Title", "Content"], supp_tables, widths=[4.2, 6.4, 7.0])

    doc.add_heading("八、正文主图与关键补充图预览", level=1)
    add_p(doc, "以下页面把必须进入正文的主图和有价值的补充图放入同一 Word 文件，便于后续逐项精修。正文图建议最终统一导出为 TIFF 330 DPI 以上；当前 Word 使用 PNG 预览以保证可快速检查。")

    set_landscape(doc)
    add_figure(
        doc,
        "framework",
        "Figure 1. MES-enabled issue-driven hierarchical batch diagnosis framework.",
        "正文必需。用于总览 issue identification、batch linkage、layer-wise screening、joint modeling、temporal adjustment、causal-informed decomposition 和 graph evidence scoring。",
        width_cm=25.8,
        break_before=False,
    )
    add_figure(
        doc,
        "data",
        "Figure 2. Real-world manufacturing dataset and batch linkage.",
        "正文必需。用于展示五类数据实体、记录数和层级批次关联。",
        width_cm=25.5,
    )
    add_figure(
        doc,
        "d2_disintegration",
        "Figure 3A. Finished-product disintegration time as the sentinel quality issue.",
        "正文必需。展示 0.8 g 成品崩解时限的时间变化和 >10 min 问题阈值。",
        width_cm=25.5,
    )
    add_two_figures(
        doc,
        "d2_weight",
        "Figure 3B. Coated tablet weight by disintegration issue group.",
        "d2_content",
        "Figure 3C. Active content by disintegration issue group.",
    )
    add_figure(
        doc,
        "d3_rank",
        "Figure 4A. Tablet MES feature screening rank.",
        "正文必需。成品 MES 是诊断增益最大的过程层。",
        width_cm=20.5,
    )
    add_two_figures(
        doc,
        "d4_rank",
        "Figure 4B. Extract-powder feature screening rank.",
        "d6_rank",
        "Figure 4C. Chenpi feature screening rank.",
    )
    add_figure(
        doc,
        "d7_rank",
        "Figure 4D. Chinese yam powder MES feature screening rank.",
        "正文必需。与 Figure 4A-C 共同构成层级筛选主图。",
        width_cm=21.5,
    )
    add_two_figures(
        doc,
        "model_auc",
        "Figure 5A. AUC and PR-AUC of hierarchical diagnostic models.",
        "model_brier",
        "Figure 5B. Brier score of hierarchical diagnostic models.",
    )
    add_figure(
        doc,
        "time_series",
        "Figure 6A. Finished-product disintegration time-series pattern.",
        "正文必需。用于说明异常并非均匀出现，而是具有时间窗口聚集。",
        width_cm=23.0,
    )
    add_two_figures(
        doc,
        "monthly_issue",
        "Figure 6B. Monthly disintegration issue rate.",
        "causal_chenpi",
        "Figure 6C. Chenpi source path attenuation after pathway adjustment.",
    )
    add_figure(
        doc,
        "evidence_chain",
        "Figure 7A. Evidence chain from finished-product issue to upstream risk-priority output.",
        "正文必需。展示从成品崩解问题到上游风险优先级输出的完整证据链。",
        width_cm=25.5,
    )
    add_two_figures(
        doc,
        "graph_top",
        "Figure 7B. Graph-based evidence score of top risk-priority nodes.",
        "graph_landscape",
        "Figure 7C. Time-adjusted batch evidence landscape.",
    )

    set_portrait(doc)
    doc.add_heading("九、建议进入支撑材料的关键图预览", level=1)
    add_p(doc, "以下图不建议放入正文主体，但对支撑完整性有价值，应作为 Supplementary Figures 或 Supplementary Data 保留。")
    set_landscape(doc)
    add_two_figures(doc, "d3_key", "Supplementary Figure S2. Key MES features vs disintegration time.", "d4_key", "Supplementary Figure S3. Key extract-powder features vs disintegration.", break_before=False)
    add_two_figures(doc, "d6_key", "Supplementary Figure S4. Key Chenpi features vs linked issue ratio.", "d7_key", "Supplementary Figure S5. Key Chinese yam powder features vs linked issue ratio.")
    add_two_figures(doc, "model_coef", "Supplementary Figure S6. Clean elastic-net coefficients.", "time_model", "Supplementary Figure S7. Time-adjusted model performance.")
    add_two_figures(doc, "causal_exposure", "Supplementary Figure S8. Upstream issue rate and abnormal-window exposure.", "causal_bootstrap", "Supplementary Figure S9. Bootstrap batch-level effects.")
    add_figure(doc, "graph_distribution", "Supplementary Figure S10. Layer-wise graph evidence score distribution.", "用于支撑图证据评分模块的层级分布。", width_cm=20.0)

    doc.save(str(OUT_DOCX))
    return doc


def export_pdf():
    if not SOFFICE.exists():
        return []
    if OUT_PDF.exists():
        OUT_PDF.unlink()
    subprocess.run(
        [
            str(SOFFICE),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(WRITE_DIR),
            str(OUT_DOCX),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    return render_pdf_preview(OUT_PDF)


def main():
    build_doc()
    pages = export_pdf()
    print("DOCX:", OUT_DOCX)
    if OUT_PDF.exists():
        print("PDF:", OUT_PDF)
    print("PREVIEW_PAGES:", len(pages))
    for page in pages[:20]:
        print("PREVIEW:", page)


if __name__ == "__main__":
    main()
