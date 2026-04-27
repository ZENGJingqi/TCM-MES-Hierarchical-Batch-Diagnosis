from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path.cwd()
IN_DOC = ROOT / "manuscript_update_in.docx"
OUT_ASCII = ROOT / "manuscript_update_out.docx"
OUT_DOC = ROOT / "论文写作" / "论文初稿-0426_完善版.docx"


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10.5)


def main() -> None:
    doc = Document(IN_DOC)
    replacements = {
        "This study used real-world manufacturing data from Jianwei Xiaoshi tablets and integrated 5,975 valid batch-level records": (
            "This study used real-world manufacturing data from Jianwei Xiaoshi tablets and integrated 5,975 valid batch-level records "
            "from quality-control testing and MES production records between August 2024 and March 2026. The dataset included five "
            "batch-level data entities: 3,728 Jianwei Xiaoshi tablet quality-testing records, 1,243 Jianwei Xiaoshi tablet MES production "
            "records, 618 Jianwei Xiaoshi extract-powder quality-testing records, 44 Chenpi quality-testing records, and 342 Chinese yam "
            "powder MES production records. Figure 1 summarizes the dataset as a manufacturing-oriented batch-linkage map rather than "
            "as isolated tables. In this map, linked denotes records or batches in the current layer that were connected forward to the "
            "next manufacturing layer, whereas traced denotes downstream records or batches that could be traced back to the current layer."
        ),
        "Figure 1. Overview of the real-world batch-level manufacturing dataset and batch-linkage structure.": (
            "Figure 1. Overview of the real-world batch-level manufacturing dataset and batch-linkage structure. The figure summarizes "
            "five data entities from Jianwei Xiaoshi tablet manufacturing: Chenpi quality testing, Chinese yam powder MES production "
            "records, Jianwei Xiaoshi extract-powder testing, Jianwei Xiaoshi tablet MES production records, and Jianwei Xiaoshi tablet "
            "quality testing. linked indicates current-layer records or batches connected forward to the next manufacturing layer; traced "
            "indicates downstream records or batches traceable back to the current layer."
        ),
        "Among the 3,728 Jianwei Xiaoshi tablet quality-testing records, 908 records were linked": (
            "The batch-linkage map showed that 41 of 44 Chenpi testing batches were linked forward to Jianwei Xiaoshi extract-powder "
            "batches, covering 578 downstream extract-powder batches that were traceable back to the Chenpi layer. For Chinese yam powder, "
            "176 of 342 batches were linked forward to tablet MES production records, and 1,159 downstream tablet MES records were traceable "
            "back to these Chinese yam powder batches. For Jianwei Xiaoshi extract powder, 273 of 618 extract-powder batches were linked "
            "forward to tablet MES production records, and 933 downstream tablet MES records were traceable back to the extract-powder layer. "
            "Finally, 908 Jianwei Xiaoshi tablet quality-testing records were linked to tablet MES production records through finished-product "
            "batch identifiers. This structure provided the data foundation for tracing finished-product disintegration issues to upstream "
            "material and manufacturing-process records."
        ),
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for prefix, replacement in replacements.items():
            if text.startswith(prefix):
                set_paragraph_text(paragraph, replacement)
                break
    doc.save(OUT_ASCII)
    shutil.copyfile(OUT_ASCII, OUT_DOC)
    print(OUT_DOC)


if __name__ == "__main__":
    main()
