from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_DIR = Path(r"D:\博士后文件\论文撰写\华润江中大数据分析")
OUT = PROJECT_DIR / "论文写作" / "论文框架_正文与支撑材料_投稿级设计稿.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_doc_fonts(doc):
    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 18, "222222"),
        ("Heading 1", 15, "222222"),
        ("Heading 2", 12.5, "222222"),
        ("Heading 3", 11.5, "222222"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_para(doc, text, style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        rest = text[len(bold_lead):]
        r2 = p.add_run(rest)
        r2.font.name = "Arial"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    else:
        r = p.add_run(text)
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, bold=True, size=font_size)
        set_cell_shading(hdr_cells[i], "EDEDED")
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            hdr_cells[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_section_heading(doc, number, title):
    doc.add_heading(f"{number}. {title}", level=1)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    set_doc_fonts(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("投稿级论文框架设计稿")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("MES-enabled issue-driven hierarchical batch diagnosis for traditional Chinese medicine manufacturing")
    r.italic = True
    r.font.size = Pt(11)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    add_para(doc, "用途：本文件用于固定论文投稿级结构、正文图表体系、支撑材料体系和写作边界。后续正式写作、图表精修和补充材料整理均以此为主线。")

    add_section_heading(doc, "一", "拟定题目与论文定位")
    add_para(doc, "英文题目：Tracing finished-product quality issues to upstream material and process risks in traditional Chinese medicine manufacturing using MES-enabled hierarchical batch diagnosis", bold_lead="英文题目：")
    add_para(doc, "中文题目：基于 MES 的中药制造成品质量问题分层批次诊断：从成品崩解异常到上游物料与过程风险追溯", bold_lead="中文题目：")
    add_para(doc, "论文定位：本文不是单纯的数据分析论文，也不是证明某一个原料、产地或供应商直接导致成品崩解异常的强因果论文。本文应定位为一个真实世界中药制造质量诊断框架研究，即将成品质量问题作为入口，整合 QC、MES 和批次追溯数据，建立多层级证据链，并输出上游物料与过程批次的风险优先级。", bold_lead="论文定位：")
    add_para(doc, "核心问题：真实世界中药制造中，成品质量异常被发现后，如何系统追溯到上游物料与过程风险，并形成可解释、可复核、可用于现场排查的批次证据链。", bold_lead="核心问题：")

    add_section_heading(doc, "二", "核心创新点")
    add_bullets(doc, [
        "提出 finished-product-issue-driven hierarchical batch diagnosis 框架，将成品质量异常转化为分层批次诊断问题。",
        "构建覆盖成品质量、成品 MES、浸膏粉质量、陈皮质量和山药粉 MES 的真实世界中药制造数据集。",
        "将批次追溯关系作为核心数据结构，而不仅仅是数据清洗过程中的辅助变量。",
        "通过 layer-wise screening、hierarchical joint modeling、temporal-window adjustment 和 causal-informed path decomposition，避免把时间窗口或来源信号过度解释为直接原因。",
        "提出 graph-based batch evidence scoring module，将多层级证据转化为可操作的上游批次 risk-priority ranking。",
    ])

    add_section_heading(doc, "三", "正文主线与 Results 结构")
    result_rows = [
        ["Results 1", "真实世界多层级数据集与批次追溯结构", "说明 2024 年 8 月至 2026 年 3 月期间 5,975 条有效批次级记录，覆盖成品理化检测、成品 MES、浸膏粉检测、陈皮检测和山药粉 MES。重点强调批次如何从成品向上游追溯。", "Figure 1; Table 1"],
        ["Results 2", "成品崩解异常作为问题驱动入口", "以成品崩解时限 >10 min 定义 sentinel finished-product issue，说明异常具有时间聚集性，适合作为框架入口。片重和含量作为辅助质量属性，不作为全文主问题。", "Figure 2"],
        ["Results 3", "逐层筛选候选质量与过程信号", "在成品 MES、浸膏粉、陈皮和山药粉层级分别筛选与崩解异常相关的候选信号。该节定位为生成候选证据，不做因果定论。", "Figure 3"],
        ["Results 4", "多层级联合建模证明诊断增益", "比较 Baseline、Baseline + MES、Full core hierarchy 等模型。重点写 MES 和层级批次信息显著提高回顾性诊断能力。", "Figure 4; Table 2"],
        ["Results 5", "时间窗口校正与因果启发式路径分解", "识别 abnormal temporal windows，说明时间窗口是关键系统状态信号。陈皮来源或产地信号经路径调整后不支持直接因果解释。", "Figure 5"],
        ["Results 6", "图证据评分与上游风险优先级", "将批次追溯关系构造成 batch-evidence graph，整合 time-adjusted residual、graph propagation、trace support 和 quality/process deviation，输出上游批次 risk-priority score。", "Figure 6; Table 3"],
    ]
    add_table(doc, ["章节", "主题", "正文内容", "推荐图表"], result_rows, font_size=8.2)

    add_section_heading(doc, "四", "正文主图与主表")
    figure_rows = [
        ["Figure 1", "Real-world hierarchical manufacturing dataset and batch linkage", "展示五类数据实体、记录数、时间范围和 linked/traced 批次关系。", "正文开篇，建立数据基础和中药多层级制造特征。"],
        ["Figure 2", "Finished-product disintegration as the sentinel quality issue", "展示成品崩解时限时间趋势、>10 min 阈值、异常窗口和月度 issue rate。", "证明为什么全文聚焦崩解时限。"],
        ["Figure 3", "Layer-wise screening of candidate process and material-quality signals", "四个 panel 展示成品 MES、浸膏粉、陈皮和山药粉层级的 FDR screening rank。", "连接成品问题与各层级候选证据。"],
        ["Figure 4", "Hierarchical joint modeling and diagnostic performance", "展示 AUC、PR-AUC、Brier 和核心变量贡献。", "证明 MES 和多层级数据的诊断增益。"],
        ["Figure 5", "Temporal-window adjustment and causal-informed path decomposition", "展示异常窗口、上游 batch exposure 和路径衰减/稳健性。", "避免将时间窗口或来源信号误写成直接原因。"],
        ["Figure 6", "Graph-based batch evidence scoring and upstream risk-priority output", "展示 batch-evidence graph、scoring module 和 top risk-priority batches。", "作为框架落地和决策支持模块。"],
    ]
    add_table(doc, ["图号", "建议标题", "内容", "正文作用"], figure_rows, font_size=8.0)

    table_rows = [
        ["Table 1", "Real-world manufacturing dataset and batch-linkage summary", "五类数据实体、记录数、主要批次标识、关联方式、linked/traced 数量。"],
        ["Table 2", "Hierarchical diagnostic model performance", "Baseline、Baseline + MES、Full hierarchy、abnormal-window models 的 AUC、PR-AUC、Brier 和交叉验证概况。"],
        ["Table 3", "Top-ranked upstream batches from graph-based evidence scoring", "上游层级、批号、linked records、issue records、time-adjusted residual、risk-priority score、support flag 和 interpretation。"],
    ]
    add_table(doc, ["表号", "建议标题", "内容"], table_rows, font_size=8.2)

    add_section_heading(doc, "五", "支撑材料设计")
    supp_fig_rows = [
        ["Supplementary Figure S1", "Finished-product quality attributes beyond disintegration", "成品片重、含量和崩解时限完整描述图。"],
        ["Supplementary Figure S2", "Finished-product MES feature screening", "成品 MES 全部有效过程参数筛选结果。"],
        ["Supplementary Figure S3", "Extract-powder quality feature screening", "浸膏粉理化属性与成品崩解异常关联。"],
        ["Supplementary Figure S4", "Chenpi quality and source-related analyses", "陈皮质量属性、产地/来源模式和与下游崩解关联。"],
        ["Supplementary Figure S5", "Chinese yam powder MES feature screening", "山药粉 MES 有效过程参数与成品崩解异常关联。"],
        ["Supplementary Figure S6", "Elastic-net coefficients for hierarchical models", "联合模型核心系数。"],
        ["Supplementary Figure S7", "XGBoost SHAP feature importance", "非线性模型变量重要性。"],
        ["Supplementary Figure S8", "Missingness of model variables by data layer", "建模变量缺失情况。"],
        ["Supplementary Figure S9", "Monthly and seasonal sensitivity analyses", "月度趋势、季节敏感性、异常窗口排除分析。"],
        ["Supplementary Figure S10", "Causal-informed path model performance", "因果启发式路径模型性能。"],
        ["Supplementary Figure S11", "Chenpi source path attenuation", "陈皮来源或产地项在路径调整后的衰减。"],
        ["Supplementary Figure S12", "Layer-wise graph-based risk-priority score distribution", "图证据评分在不同上游层级中的分布。"],
        ["Supplementary Figure S13", "Time-adjusted batch evidence landscape", "上游批次 linked records 与 time-adjusted residual 的关系。"],
    ]
    add_table(doc, ["补充图", "标题", "内容"], supp_fig_rows, font_size=7.6)

    supp_table_rows = [
        ["Supplementary Table S1", "Data dictionary and variable-level completeness", "所有定稿英文数据表的字段说明、数据类型、缺失值、数值范围、类别变量情况。"],
        ["Supplementary Table S2", "Batch linkage and traceability summary", "各层级 linked/traced 批次数、记录数和关联比例。"],
        ["Supplementary Table S3", "Finished-product quality descriptive statistics", "成品理化数据描述性统计和 >10 min 分组比较。"],
        ["Supplementary Table S4", "Layer-wise feature screening results", "成品 MES、浸膏粉、陈皮、山药粉所有有效指标的 P 值、FDR、效应方向。"],
        ["Supplementary Table S5", "Full hierarchical model results", "所有模型性能、交叉验证、时间拆分验证、核心变量结果。"],
        ["Supplementary Table S6", "Temporal and causal-informed analyses", "异常窗口、季节性、路径衰减、bootstrap batch-level effects。"],
        ["Supplementary Table S7", "Graph-based batch evidence scoring results", "所有上游批次的 graph evidence score、risk-priority score 和解释标签。"],
        ["Supplementary Data 1", "Analysis-ready English dataset", "定稿英文数据文件打包。"],
        ["Supplementary Data 2", "Analysis code", "R 分析代码和图表生成代码。"],
    ]
    add_table(doc, ["支撑材料", "标题", "内容"], supp_table_rows, font_size=7.8)

    add_section_heading(doc, "六", "Methods 写作框架")
    method_rows = [
        ["Data curation and standardization", "说明中文原始数据如何标准化为英文 analysis-ready datasets，批号转字符串，重复批号处理，多值字段原则，D7 缺失值均值插补规则。"],
        ["Batch linkage construction", "说明成品质量、成品 MES、浸膏粉、陈皮和山药粉之间的批次标识与追溯逻辑。"],
        ["Definition of sentinel issue", "说明崩解时限 >10 min 的定义依据和分析范围。"],
        ["Layer-wise statistical screening", "Spearman 相关、Wilcoxon 检验、BH-FDR 校正、linked issue ratio 分析。"],
        ["Hierarchical diagnostic modeling", "说明模型层级、特征纳入原则、缺失处理、交叉验证、AUC、PR-AUC、Brier。"],
        ["Temporal-window adjustment", "说明异常窗口识别、月度趋势、季节敏感性和时间校正模型。"],
        ["Causal-informed path decomposition", "说明路径假设、bootstrap batch-level robustness 和 source-term attenuation，强调不是随机因果推断。"],
        ["Graph-based batch evidence scoring", "说明 graph evidence score、risk-priority score、graph propagation、trace support 和 quality/process deviation 的计算。"],
    ]
    add_table(doc, ["方法模块", "应写内容"], method_rows, font_size=8.2)

    add_section_heading(doc, "七", "关键结果的投稿级表述")
    add_bullets(doc, [
        "本研究构建了一个真实世界中药制造多层级数据集，而不是单一质检数据集。",
        "成品崩解时限 >10 min 可作为 sentinel finished-product issue，引导后续分层诊断。",
        "MES 数据提供了最大的回顾性诊断增益，说明生产过程记录对成品质量问题解释具有核心价值。",
        "异常时间窗口是重要系统状态信号，因此上游原料或来源解释必须经过 temporal-window adjustment。",
        "陈皮来源相关信号不应写成直接原因，应写为在时间和路径调整后不支持稳定直接归因。",
        "Graph-based batch evidence scoring 将统计证据和追溯结构转化为上游批次优先级，是本文框架落地的关键模块。",
    ])

    add_section_heading(doc, "八", "写作边界与风险控制")
    add_table(
        doc,
        ["可以写", "不要写"],
        [
            ["MES-enabled hierarchical batch diagnosis framework。", "某个陈皮产地或某个供应商被证明导致崩解异常。"],
            ["Finished-product-issue-driven retrospective diagnosis。", "模型已经可以稳定预测未来所有异常。"],
            ["Causal-informed path decomposition。", "这是随机因果推断或强因果证明。"],
            ["Graph-based batch evidence scoring for upstream risk prioritization。", "risk-priority score 是因果效应大小或显著性 P 值。"],
            ["Prioritized upstream batches for focused investigation。", "直接给出现场责任归因。"],
        ],
        font_size=8.2,
    )

    add_section_heading(doc, "九", "当前完成度判断")
    add_para(doc, "当前结果已经足够支撑一篇框架型真实世界制造数据论文。最强的部分是：数据结构清楚、成品问题明确、MES 诊断增益明显、时间窗口解释充分、因果写法边界明确、图证据评分模块形成了可操作输出。", bold_lead="当前结果已经足够支撑一篇框架型真实世界制造数据论文。")
    add_para(doc, "后续最重要的工作不是继续堆更多单变量分析，而是把主图、Results 文字、Methods 和 Supplementary materials 按上述框架收束。正文必须围绕一个问题展开：如何将成品崩解异常转化为上游物料与过程风险的多层级批次诊断证据链。")

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("附：建议正文摘要草稿", level=1)
    add_para(doc, "真实世界中药制造通常涉及多饮片、多中间体、多工序和多批次追溯关系，常规成品质量控制检测虽能识别异常批次，但难以系统追溯上游物料和过程风险。本研究以健胃消食片真实世界生产数据为对象，整合 2024 年 8 月至 2026 年 3 月期间来自质量控制检测和 MES 生产记录的 5,975 条有效批次级记录，构建覆盖成品质量、成品生产过程、中间体质量、原料质量和过程物料生产的多层级制造数据集。我们以成品崩解时限 >10 min 作为 sentinel finished-product quality issue，提出 MES-enabled issue-driven hierarchical batch diagnosis 框架，依次开展批次追溯、逐层特征筛选、多层级联合建模、时间窗口校正、因果启发式路径分解和图证据评分。结果显示，加入 MES 和层级批次信息显著提高了成品崩解异常的回顾性诊断能力；异常时间窗口是关键系统状态信号，直接将上游来源信号解释为因果来源并不稳妥。进一步构建的 graph-based batch evidence scoring module 将 time-adjusted residual、graph propagation、trace support 和 quality/process deviation 整合为上游批次 risk-priority score，为现场 focused investigation 提供可解释的批次优先级。本研究提供了一种将真实世界 MES 和质量控制数据转化为中药制造质量问题诊断证据链的可推广框架。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
